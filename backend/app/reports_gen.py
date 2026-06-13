"""Synchronous report generation: gather org posture, render to PDF / Word / Excel.

Reports
  readiness        — control-by-control posture with evidence status
  gap              — controls not yet compliant, worst-risk first
  executive        — one-page KPI + maturity summary
  recommendations  — Top 10 priorities, quick wins, then the full list
  evidence         — evidence register (Excel only)

PDF via WeasyPrint, Word via python-docx, Excel via openpyxl. All produced in
memory and returned as (bytes, mime, filename).
"""
from __future__ import annotations

import io
import os
import re
import zipfile
from datetime import datetime, timezone

from app.core.config import settings

from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.tenant import Tenant
from app.models.framework import Framework, Control
from app.models.evidence import OrgControl, EvidenceItem
from app.models.self_assessment import SelfAssessment
from app.models.recommendation import Recommendation
from app.api.controls import _display_status, AUDIT_LABEL
from app.core import recommendations as rec_engine
from app.report_assets import EVA_LOGO_DATAURI

BRAND = "#1A8FD1"
EV_STATUS_LABEL = {
    "accepted": "Accepted", "rejected": "Rejected", "needs_more": "Needs more",
    "eva_pending": "Pending EVA", "client_submitted": "Submitted", "draft": "Draft",
}


# ──────────────────────────── data ────────────────────────────
async def gather(db: AsyncSession, org_id) -> dict:
    tenant = (await db.execute(select(Tenant).where(Tenant.id == org_id))).scalar_one_or_none()
    client_name = tenant.name if tenant else "Client"

    fw_map = {fid: name for fid, name in (await db.execute(select(Framework.id, Framework.name))).all()}
    sub_fws = list((await db.execute(
        select(Control.framework_id).join(OrgControl, OrgControl.control_id == Control.id)
        .where(OrgControl.org_id == org_id).distinct()
    )).scalars().all())
    controls = (await db.execute(
        select(Control).where(Control.framework_id.in_(sub_fws)).order_by(Control.sort_order)
    )).scalars().all() if sub_fws else []

    ocs = {oc.control_id: oc for oc in (await db.execute(
        select(OrgControl).where(OrgControl.org_id == org_id)
    )).scalars().all()}
    ev_counts = {oc_id: n for oc_id, n in (await db.execute(
        select(EvidenceItem.org_control_id, func.count(EvidenceItem.id))
        .where(EvidenceItem.org_id == org_id).group_by(EvidenceItem.org_control_id)
    )).all()}
    answered = {cid for cid, ans in (await db.execute(
        select(SelfAssessment.control_id, SelfAssessment.answers).where(SelfAssessment.org_id == org_id)
    )).all() if ans}

    rows = []
    counts = {"not_started": 0, "in_progress": 0, "compliant": 0, "non_compliant": 0, "not_applicable": 0}
    for c in controls:
        oc = ocs.get(c.id)
        evc = ev_counts.get(oc.id, 0) if oc else 0
        ds = _display_status(oc, evc, c.id in answered)
        counts[ds] = counts.get(ds, 0) + 1
        rows.append({
            "ref": c.ref, "title": c.title, "domain": c.domain or "General",
            "framework": fw_map.get(c.framework_id, "—"),
            "risk": (c.risk_rating.value if c.risk_rating else "low"),
            "priority": (c.priority.value if c.priority else "low"),
            "status": ds, "status_label": AUDIT_LABEL.get(ds, ds),
            "coverage": (oc.coverage_pct if oc else 0),
            "evidence": evc,
        })

    total = len(rows) or 1
    compliant = counts.get("compliant", 0)
    compliance_pct = round(compliant / total * 100)

    # Evidence register
    ev_rows = (await db.execute(
        select(EvidenceItem, Control.ref, Control.title, Control.framework_id)
        .join(OrgControl, OrgControl.id == EvidenceItem.org_control_id)
        .join(Control, Control.id == OrgControl.control_id)
        .where(EvidenceItem.org_id == org_id)
        .order_by(Control.framework_id, Control.ref)
    )).all()
    evidence = [{
        "framework": fw_map.get(fwid, "Framework"),
        "ctrl_ref": ref, "ctrl_title": title, "title": ev.title,
        "status": EV_STATUS_LABEL.get(ev.status.value, ev.status.value),
        "file_name": ev.file_name or "—",
        "file_key": ev.file_key or "",
        "note": ev.review_note or "",
        "date": ev.created_at.strftime("%Y-%m-%d") if ev.created_at else "—",
    } for ev, ref, title, fwid in ev_rows]

    # Recommendations
    rec_rows = (await db.execute(
        select(Recommendation, Control)
        .join(Control, Control.id == Recommendation.control_id)
        .where(Recommendation.org_id == org_id)
    )).all()
    recs = []
    for r, c in rec_rows:
        gap = (r.target_level - r.current_level) if (r.target_level is not None and r.current_level is not None) else None
        recs.append({
            "control_ref": c.ref, "control_title": c.title, "domain": c.domain or "General",
            "title": r.title, "text": r.text, "effort": r.effort, "impact": r.impact,
            "status": r.status, "is_top10": bool(r.is_top10),
            "quick_win": rec_engine.is_quick_win(r.effort, r.impact),
            "priority": rec_engine.priority_score(r.effort, r.impact, gap, rec_engine._risk_value(c.risk_rating)),
        })
    recs.sort(key=lambda x: x["priority"], reverse=True)
    active = [r for r in recs if r["status"] not in ("done", "dismissed")]
    pinned = [r for r in active if r["is_top10"]]
    top10 = pinned if pinned else active[:10]
    quick_wins = [r for r in active if r["quick_win"]]

    # Maturity overall
    try:
        from app.api.maturity import _overall
        maturity = await _overall(db, org_id)
    except Exception:
        maturity = {"perceived": None, "assessed": 0, "target": 0, "gap": None, "has_data": False}

    return {
        "client": client_name,
        "generated_at": datetime.now(timezone.utc).strftime("%B %d, %Y"),
        "frameworks": sorted({r["framework"] for r in rows}),
        "controls": rows,
        "counts": counts,
        "total": len(rows),
        "compliance_pct": compliance_pct,
        "evidence": evidence,
        "recommendations": recs,
        "top10": top10,
        "quick_wins": quick_wins,
        "maturity": maturity,
    }


REPORT_TITLES = {
    "readiness": "Audit Readiness Report",
    "gap": "Gap Analysis",
    "executive": "Executive Summary",
    "recommendations": "Recommendations & Top 10 Priorities",
    "evidence": "Evidence Register",
}


# ──────────────────────────── HTML / PDF ────────────────────────────
def _esc(s) -> str:
    return (str(s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


_BADGE = {
    "compliant": "#16A34A", "non_compliant": "#DC2626", "in_progress": "#D97706",
    "not_started": "#64748B", "not_applicable": "#64748B",
    "high": "#DC2626", "critical": "#DC2626", "medium": "#D97706", "low": "#16A34A",
}


def _html_shell(title: str, d: dict, body: str) -> str:
    year = datetime.now(timezone.utc).year
    footer = f"Confidential — EVA Technologies {year}"
    return f"""<!doctype html><html><head><meta charset="utf-8"><style>
    @page {{ size: A4; margin: 1.7cm 1.6cm 2cm;
      @bottom-left {{ content: "{footer}"; font-size: 8px; color: #94A3B8; }}
      @bottom-right {{ content: "Page " counter(page) " of " counter(pages); font-size: 8px; color: #94A3B8; }}
    }}
    body {{ font-family: 'Helvetica Neue', Arial, sans-serif; color: #0F172A; font-size: 11px; line-height: 1.5; }}
    .brandbar {{ display: flex; align-items: center; gap: 10px; margin-bottom: 14px; padding-bottom: 12px; border-bottom: 1px solid #E2E8F0; }}
    .brandbar img {{ height: 42px; }}
    .brandbar .tag {{ margin-left: auto; font-size: 9px; letter-spacing: .08em; text-transform: uppercase; color: #94A3B8; }}
    .cover {{ border-left: 6px solid {BRAND}; padding: 4px 0 4px 14px; margin-bottom: 22px; }}
    .cover h1 {{ font-size: 22px; margin: 0 0 4px; color: {BRAND}; }}
    .cover .meta {{ color: #475569; font-size: 11px; }}
    h2 {{ font-size: 14px; color: {BRAND}; border-bottom: 2px solid #E2E8F0; padding-bottom: 4px; margin: 22px 0 10px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 8px 0; }}
    th {{ background: #F1F5F9; text-align: left; padding: 6px 8px; font-size: 9px; text-transform: uppercase; letter-spacing: .04em; color: #475569; border-bottom: 1px solid #CBD5E1; }}
    td {{ padding: 6px 8px; border-bottom: 1px solid #E2E8F0; vertical-align: top; }}
    .kpis {{ display: flex; gap: 12px; margin: 12px 0; }}
    .kpi {{ flex: 1; border: 1px solid #E2E8F0; border-radius: 8px; padding: 12px; }}
    .kpi .v {{ font-size: 24px; font-weight: 700; color: {BRAND}; }}
    .kpi .l {{ font-size: 10px; color: #64748B; text-transform: uppercase; letter-spacing: .04em; }}
    .badge {{ display: inline-block; white-space: nowrap; font-size: 9px; font-weight: 700; padding: 2px 7px; border-radius: 10px; color: #fff; }}
    td .badge {{ line-height: 1.4; }}
    .top {{ border: 1px solid #FBBF24; background: #FFFBEB; border-radius: 8px; padding: 10px 12px; margin-bottom: 8px; }}
    .top .rank {{ font-weight: 800; color: #B45309; }}
    .qw {{ color: #16A34A; font-weight: 700; }}
    .muted {{ color: #64748B; }}
    </style></head><body>
    <div class="brandbar">
      <img src="{EVA_LOGO_DATAURI}" alt="EVA Technologies" />
      <span class="tag">Confidential</span>
    </div>
    <div class="cover"><h1>{_esc(title)}</h1>
      <div class="meta">{_esc(d['client'])} &nbsp;·&nbsp; {_esc(', '.join(d['frameworks']) or '—')} &nbsp;·&nbsp; {_esc(d['generated_at'])}</div>
    </div>
    {body}
    </body></html>"""


def _badge(text: str, key: str) -> str:
    return f'<span class="badge" style="background:{_BADGE.get(key, "#64748B")}">{_esc(text)}</span>'


def _body_readiness(d: dict) -> str:
    c = d["counts"]
    rows = "".join(
        f"<tr><td><b>{_esc(r['ref'])}</b></td><td>{_esc(r['title'])}</td><td>{_esc(r['domain'])}</td>"
        f"<td>{_badge(r['status_label'], r['status'])}</td><td>{r['coverage']}%</td><td>{r['evidence']}</td></tr>"
        for r in d["controls"]
    )
    return f"""<div class="kpis">
      <div class="kpi"><div class="v">{d['compliance_pct']}%</div><div class="l">Compliant</div></div>
      <div class="kpi"><div class="v">{c.get('compliant',0)}</div><div class="l">Compliant</div></div>
      <div class="kpi"><div class="v">{c.get('in_progress',0)}</div><div class="l">In progress</div></div>
      <div class="kpi"><div class="v">{c.get('not_started',0)}</div><div class="l">Not started</div></div>
      <div class="kpi"><div class="v">{d['total']}</div><div class="l">Total controls</div></div>
    </div>
    <h2>Control-by-control posture</h2>
    <table><thead><tr><th>Ref</th><th>Control</th><th>Domain</th><th>Status</th><th>Coverage</th><th>Evidence</th></tr></thead>
    <tbody>{rows}</tbody></table>"""


def _body_gap(d: dict) -> str:
    risk_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "informational": 4}
    gaps = [r for r in d["controls"] if r["status"] in ("not_started", "in_progress", "non_compliant")]
    gaps.sort(key=lambda r: (risk_order.get(r["risk"], 5), r["coverage"]))
    rows = "".join(
        f"<tr><td><b>{_esc(r['ref'])}</b></td><td>{_esc(r['title'])}</td>"
        f"<td>{_badge(r['risk'].title(), r['risk'])}</td><td>{_badge(r['status_label'], r['status'])}</td>"
        f"<td>{r['coverage']}%</td></tr>"
        for r in gaps
    )
    return f"""<p class="muted">{len(gaps)} of {d['total']} controls are not yet compliant, ordered by risk.</p>
    <table><thead><tr><th>Ref</th><th>Control</th><th>Risk</th><th>Status</th><th>Coverage</th></tr></thead>
    <tbody>{rows or '<tr><td colspan=5 class=muted>No open gaps — all controls compliant.</td></tr>'}</tbody></table>"""


def _body_executive(d: dict) -> str:
    m = d["maturity"]
    c = d["counts"]
    mat = (f"<div class='kpi'><div class='v'>{m['assessed']}/5</div><div class='l'>Assessed maturity</div></div>"
           f"<div class='kpi'><div class='v'>{m['perceived'] if m.get('perceived') is not None else '—'}</div><div class='l'>Perceived maturity</div></div>") if m.get("has_data") else ""
    return f"""<div class="kpis">
      <div class="kpi"><div class="v">{d['compliance_pct']}%</div><div class="l">Compliance</div></div>
      <div class="kpi"><div class="v">{c.get('compliant',0)}/{d['total']}</div><div class="l">Controls compliant</div></div>
      {mat}
    </div>
    <h2>Posture at a glance</h2>
    <table><thead><tr><th>Status</th><th>Controls</th></tr></thead><tbody>
      <tr><td>{_badge('Compliant','compliant')}</td><td>{c.get('compliant',0)}</td></tr>
      <tr><td>{_badge('In Progress','in_progress')}</td><td>{c.get('in_progress',0)}</td></tr>
      <tr><td>{_badge('Non-Compliant','non_compliant')}</td><td>{c.get('non_compliant',0)}</td></tr>
      <tr><td>{_badge('Not Started','not_started')}</td><td>{c.get('not_started',0)}</td></tr>
    </tbody></table>
    <h2>Top priorities</h2>
    <ol>{''.join(f"<li><b>{_esc(r['control_ref'])}</b> — {_esc(r['title'])}</li>" for r in d['top10'][:5]) or '<li class=muted>No recommendations generated yet.</li>'}</ol>"""


def _body_recommendations(d: dict) -> str:
    def card(r, i=None):
        rank = f"<span class='rank'>#{i}</span> " if i else ""
        qw = " <span class='qw'>⚡ Quick win</span>" if r["quick_win"] else ""
        return (f"<div class='top'>{rank}<b>{_esc(r['control_ref'])}</b> — {_esc(r['title'])}{qw}"
                f"<div>{_esc(r['text'])}</div>"
                f"<div class='muted'>Impact: {_esc(r['impact'])} · Effort: {_esc(r['effort'])} · {_esc(r['domain'])}</div></div>")
    top = "".join(card(r, i + 1) for i, r in enumerate(d["top10"]))
    qw = "".join(card(r) for r in d["quick_wins"]) or "<p class='muted'>No quick wins identified.</p>"
    allrows = "".join(
        f"<tr><td><b>{_esc(r['control_ref'])}</b></td><td>{_esc(r['title'])}</td>"
        f"<td>{_esc(r['impact'])}</td><td>{_esc(r['effort'])}</td><td>{_esc(r['status'])}</td></tr>"
        for r in d["recommendations"]
    )
    return f"""<h2>Top 10 priorities</h2>{top or "<p class='muted'>No Top 10 selected yet.</p>"}
    <h2>Quick wins</h2>{qw}
    <h2>All recommendations ({len(d['recommendations'])})</h2>
    <table><thead><tr><th>Ref</th><th>Recommendation</th><th>Impact</th><th>Effort</th><th>Status</th></tr></thead>
    <tbody>{allrows or '<tr><td colspan=5 class=muted>None generated.</td></tr>'}</tbody></table>"""


_HTML_BODY = {
    "readiness": _body_readiness, "gap": _body_gap,
    "executive": _body_executive, "recommendations": _body_recommendations,
}


def render_pdf(report_type: str, d: dict) -> bytes:
    from weasyprint import HTML
    title = REPORT_TITLES.get(report_type, "Report")
    body = _HTML_BODY.get(report_type, _body_executive)(d)
    return HTML(string=_html_shell(title, d, body)).write_pdf()


# ──────────────────────────── DOCX ────────────────────────────
def render_docx(report_type: str, d: dict) -> bytes:
    import base64
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # EVA logo at the top of the document
    try:
        _logo = base64.b64decode(EVA_LOGO_DATAURI.split(",", 1)[1])
        doc.add_picture(io.BytesIO(_logo), width=Inches(1.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass

    # Confidential footer with current year on every page
    try:
        footer_p = doc.sections[0].footer.paragraphs[0]
        footer_p.text = f"Confidential — EVA Technologies {datetime.now(timezone.utc).year}"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    except Exception:
        pass

    title = REPORT_TITLES.get(report_type, "Report")
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"{d['client']} · {', '.join(d['frameworks']) or '—'} · {d['generated_at']}")

    c = d["counts"]

    def table(headers, data_rows):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, hh in enumerate(headers):
            t.rows[0].cells[i].text = hh
        for row in data_rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)

    if report_type == "readiness":
        doc.add_paragraph(f"Overall compliance: {d['compliance_pct']}%  ·  {c.get('compliant',0)} of {d['total']} controls compliant.")
        doc.add_heading("Control-by-control posture", level=1)
        table(["Ref", "Control", "Domain", "Status", "Coverage", "Evidence"],
              [[r["ref"], r["title"], r["domain"], r["status_label"], f"{r['coverage']}%", r["evidence"]] for r in d["controls"]])
    elif report_type == "gap":
        risk_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "informational": 4}
        gaps = sorted([r for r in d["controls"] if r["status"] in ("not_started", "in_progress", "non_compliant")],
                      key=lambda r: (risk_order.get(r["risk"], 5), r["coverage"]))
        doc.add_paragraph(f"{len(gaps)} of {d['total']} controls are not yet compliant, ordered by risk.")
        table(["Ref", "Control", "Risk", "Status", "Coverage"],
              [[r["ref"], r["title"], r["risk"].title(), r["status_label"], f"{r['coverage']}%"] for r in gaps])
    elif report_type == "recommendations":
        doc.add_heading("Top 10 priorities", level=1)
        for i, r in enumerate(d["top10"], 1):
            p = doc.add_paragraph()
            p.add_run(f"#{i}  {r['control_ref']} — {r['title']}").bold = True
            if r["quick_win"]:
                p.add_run("  ⚡ Quick win")
            doc.add_paragraph(r["text"])
            doc.add_paragraph(f"Impact: {r['impact']} · Effort: {r['effort']} · {r['domain']}").italic = True
        doc.add_heading(f"All recommendations ({len(d['recommendations'])})", level=1)
        table(["Ref", "Recommendation", "Impact", "Effort", "Status"],
              [[r["control_ref"], r["title"], r["impact"], r["effort"], r["status"]] for r in d["recommendations"]])
    else:  # executive
        m = d["maturity"]
        doc.add_paragraph(f"Compliance: {d['compliance_pct']}%  ·  {c.get('compliant',0)}/{d['total']} controls compliant.")
        if m.get("has_data"):
            doc.add_paragraph(f"Maturity — assessed {m['assessed']}/5, perceived {m.get('perceived', '—')}/5.")
        doc.add_heading("Posture at a glance", level=1)
        table(["Status", "Controls"],
              [["Compliant", c.get("compliant", 0)], ["In Progress", c.get("in_progress", 0)],
               ["Non-Compliant", c.get("non_compliant", 0)], ["Not Started", c.get("not_started", 0)]])
        doc.add_heading("Top priorities", level=1)
        for i, r in enumerate(d["top10"][:5], 1):
            doc.add_paragraph(f"{i}. {r['control_ref']} — {r['title']}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────── XLSX (evidence register) ────────────────────────────
def render_xlsx(report_type: str, d: dict) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Evidence Register"
    headers = ["Control", "Control Title", "Evidence", "Status", "File", "Date", "Reviewer note"]
    ws.append(headers)
    hdr_fill = PatternFill("solid", fgColor="1A8FD1")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = hdr_fill
    for e in d["evidence"]:
        ws.append([e["ctrl_ref"], e["ctrl_title"], e["title"], e["status"], e["file_name"], e["date"], e["note"]])
    widths = [14, 40, 30, 14, 28, 12, 40]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    ws.oddFooter.center.text = f"Confidential — EVA Technologies {datetime.now(timezone.utc).year}"
    ws.oddFooter.center.size = 8
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ──────────────────────────── Evidence ZIP bundle ────────────────────────────
def _safe_name(s, fallback: str = "item") -> str:
    """Filesystem-safe folder / file segment (keeps spaces, dashes, dots)."""
    s = re.sub(r"[^A-Za-z0-9._ \-]+", "_", str(s or "")).strip().strip(".")
    s = re.sub(r"\s+", " ", s).strip()
    return s or fallback


def _esc_xl(s) -> str:
    """Escape a string for use inside an Excel formula string literal."""
    return str(s or "").replace('"', '""')


def render_evidence_zip(d: dict) -> bytes:
    """A ZIP containing the Excel register plus the evidence files, foldered as
    evidences/<Framework>/<Control ref>/<file>. The register links to each folder
    with a *relative* hyperlink, so the links work wherever the ZIP is extracted."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill

    base_dir = settings.STORAGE_LOCAL_PATH

    # Plan each evidence row → its folder + (deduped) file name inside the ZIP.
    used: dict[str, set] = {}
    plans = []  # (row, rel_folder, rel_file_or_None, src_path_or_None)
    for e in d["evidence"]:
        fw = _safe_name(e.get("framework"), "Framework")
        ctrl = _safe_name(e.get("ctrl_ref"), "control")
        rel_folder = f"evidences/{fw}/{ctrl}"
        rel_file = src = None
        fk = e.get("file_key")
        if fk:
            cand_src = os.path.join(base_dir, fk)
            if os.path.exists(cand_src):
                src = cand_src
                fname = _safe_name(e.get("file_name") or os.path.basename(fk), "evidence")
                seen = used.setdefault(rel_folder, set())
                name, n = fname, 1
                while name in seen:
                    stem, dot, ext = fname.rpartition(".")
                    name = f"{stem} ({n}).{ext}" if dot else f"{fname} ({n})"
                    n += 1
                seen.add(name)
                rel_file = f"{rel_folder}/{name}"
        plans.append((e, rel_folder, rel_file, src))

    # Build the Excel register (with a relative folder hyperlink per row).
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Evidence Register"
    headers = ["Framework", "Control", "Control Title", "Evidence", "Status", "Date", "Reviewer note", "Open file", "Location (in this ZIP)"]
    ws.append(headers)
    hdr_fill = PatternFill("solid", fgColor="1A8FD1")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = hdr_fill
    link_font = Font(color="1A8FD1", underline="single")
    for e, rel_folder, rel_file, src in plans:
        # "Location" is always plain text (works everywhere — navigate the extracted ZIP).
        location = rel_file if rel_file else f"{rel_folder}/  (no file)"
        ws.append([e.get("framework", "—"), e["ctrl_ref"], e["ctrl_title"], e["title"],
                   e["status"], e["date"], e["note"], e["file_name"], location])
        # "Open file" → relative HYPERLINK to the actual document (opens the file, not a folder).
        link_cell = ws.cell(row=ws.max_row, column=8)
        if rel_file:
            link_cell.value = f'=HYPERLINK("{rel_file}","{_esc_xl(e["file_name"])}")'
            link_cell.font = link_font
    for i, w in enumerate([20, 14, 32, 26, 13, 12, 26, 26, 40], 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    ws.oddFooter.center.text = f"Confidential — EVA Technologies {datetime.now(timezone.utc).year}"
    ws.oddFooter.center.size = 8
    xbuf = io.BytesIO()
    wb.save(xbuf)

    # Assemble the ZIP.
    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("Evidence_Register.xlsx", xbuf.getvalue())
        z.writestr("README.txt",
                   "Evidence bundle\n"
                   f"Client: {d.get('client', '—')}\n"
                   f"Generated: {d.get('generated_at', '—')}\n\n"
                   "1. Extract this ZIP, then open 'Evidence_Register.xlsx'.\n"
                   "2. Each row's 'Open file' link opens that evidence document.\n"
                   "3. The 'Location' column shows the exact path inside this bundle:\n"
                   "   evidences/<Framework>/<Control>/<file>\n\n"
                   "Note: Excel for Mac is sandboxed and may block clickable local links.\n"
                   "If a link does not open, use the 'Location' path to find the file in the\n"
                   "extracted 'evidences' folder via Finder/Explorer.\n\n"
                   f"Confidential — EVA Technologies {datetime.now(timezone.utc).year}\n")
        folders_with_files = set()
        for e, rel_folder, rel_file, src in plans:
            if rel_file and src:
                try:
                    with open(src, "rb") as fh:
                        z.writestr(rel_file, fh.read())
                    folders_with_files.add(rel_folder)
                except Exception:
                    pass
        # Make sure every control's folder exists so its link always resolves.
        for e, rel_folder, rel_file, src in plans:
            if rel_folder not in folders_with_files:
                z.writestr(f"{rel_folder}/_no_files.txt",
                           "No stored evidence files for this control yet.\n")
                folders_with_files.add(rel_folder)
    return zbuf.getvalue()


# ──────────────────────────── Aggregate (MSP / global) ────────────────────────────
async def _org_posture(db: AsyncSession, org_id) -> dict:
    """Compact compliance roll-up for one client org (no evidence/recs build)."""
    sub_fws = list((await db.execute(
        select(Control.framework_id).join(OrgControl, OrgControl.control_id == Control.id)
        .where(OrgControl.org_id == org_id).distinct()
    )).scalars().all())
    controls = (await db.execute(
        select(Control).where(Control.framework_id.in_(sub_fws))
    )).scalars().all() if sub_fws else []
    ocs = {oc.control_id: oc for oc in (await db.execute(
        select(OrgControl).where(OrgControl.org_id == org_id)
    )).scalars().all()}
    ev_counts = {oc_id: n for oc_id, n in (await db.execute(
        select(EvidenceItem.org_control_id, func.count(EvidenceItem.id))
        .where(EvidenceItem.org_id == org_id).group_by(EvidenceItem.org_control_id)
    )).all()}
    answered = {cid for cid, ans in (await db.execute(
        select(SelfAssessment.control_id, SelfAssessment.answers).where(SelfAssessment.org_id == org_id)
    )).all() if ans}
    counts = {"not_started": 0, "in_progress": 0, "compliant": 0, "non_compliant": 0, "not_applicable": 0}
    for c in controls:
        oc = ocs.get(c.id)
        evc = ev_counts.get(oc.id, 0) if oc else 0
        ds = _display_status(oc, evc, c.id in answered)
        counts[ds] = counts.get(ds, 0) + 1
    total = len(controls)
    compliant = counts.get("compliant", 0)
    return {"total": total, "compliant": compliant,
            "in_progress": counts.get("in_progress", 0),
            "not_started": counts.get("not_started", 0),
            "non_compliant": counts.get("non_compliant", 0),
            "compliance_pct": round(compliant / total * 100) if total else 0}


async def gather_aggregate(db: AsyncSession, org_ids: list, scope_name: str) -> dict:
    """Posture across many client orgs (an MSP's base, or all clients)."""
    rows = []
    for oid in org_ids:
        t = (await db.execute(select(Tenant).where(Tenant.id == oid))).scalar_one_or_none()
        if not t:
            continue
        po = await _org_posture(db, oid)
        rows.append({"name": t.name, "type": t.tenant_type.value,
                     "plan": (t.plan_name or "—"), **po})
    rows.sort(key=lambda r: r["compliance_pct"])
    tot_controls = sum(r["total"] for r in rows)
    tot_compliant = sum(r["compliant"] for r in rows)
    return {
        "scope": scope_name,
        "generated_at": datetime.now(timezone.utc).strftime("%B %d, %Y"),
        "clients": rows,
        "client_count": len(rows),
        "total_controls": tot_controls,
        "total_compliant": tot_compliant,
        "avg_compliance": round(sum(r["compliance_pct"] for r in rows) / len(rows)) if rows else 0,
        "overall_compliance": round(tot_compliant / tot_controls * 100) if tot_controls else 0,
    }


def _agg_html(d: dict) -> str:
    rows = "".join(
        f"<tr><td><b>{_esc(r['name'])}</b></td><td>{_esc(r['plan'])}</td>"
        f"<td>{r['compliance_pct']}%</td><td>{r['compliant']}/{r['total']}</td>"
        f"<td>{r['in_progress']}</td><td>{r['not_started']}</td></tr>"
        for r in d["clients"]
    ) or '<tr><td colspan=6 class=muted>No clients in scope.</td></tr>'
    body = f"""<div class="kpis">
      <div class="kpi"><div class="v">{d['client_count']}</div><div class="l">Clients</div></div>
      <div class="kpi"><div class="v">{d['overall_compliance']}%</div><div class="l">Overall compliance</div></div>
      <div class="kpi"><div class="v">{d['avg_compliance']}%</div><div class="l">Average per client</div></div>
      <div class="kpi"><div class="v">{d['total_compliant']}/{d['total_controls']}</div><div class="l">Controls compliant</div></div>
    </div>
    <h2>Per-client compliance</h2>
    <table><thead><tr><th>Client</th><th>Plan</th><th>Compliance</th><th>Compliant</th><th>In progress</th><th>Not started</th></tr></thead>
    <tbody>{rows}</tbody></table>"""
    return _html_shell("Portfolio Compliance Report", {"client": d["scope"], "frameworks": [], "generated_at": d["generated_at"]}, body)


def render_aggregate_pdf(d: dict) -> bytes:
    from weasyprint import HTML
    return HTML(string=_agg_html(d)).write_pdf()


def render_aggregate_docx(d: dict) -> bytes:
    import base64
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    try:
        _logo = base64.b64decode(EVA_LOGO_DATAURI.split(",", 1)[1])
        doc.add_picture(io.BytesIO(_logo), width=Inches(1.7))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
    except Exception:
        pass
    try:
        fp = doc.sections[0].footer.paragraphs[0]
        fp.text = f"Confidential — EVA Technologies {datetime.now(timezone.utc).year}"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.size = Pt(8); fp.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    except Exception:
        pass
    doc.add_heading("Portfolio Compliance Report", level=0)
    doc.add_paragraph(f"{d['scope']} · {d['client_count']} clients · {d['generated_at']}")
    doc.add_paragraph(f"Overall compliance: {d['overall_compliance']}%  ·  average per client: {d['avg_compliance']}%  ·  {d['total_compliant']} of {d['total_controls']} controls compliant.")
    doc.add_heading("Per-client compliance", level=1)
    t = doc.add_table(rows=1, cols=6)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(["Client", "Plan", "Compliance", "Compliant", "In progress", "Not started"]):
        t.rows[0].cells[i].text = h
    for r in d["clients"]:
        cells = t.add_row().cells
        vals = [r["name"], r["plan"], f"{r['compliance_pct']}%", f"{r['compliant']}/{r['total']}", str(r["in_progress"]), str(r["not_started"])]
        for i, v in enumerate(vals):
            cells[i].text = v
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def build_aggregate_report(db: AsyncSession, org_ids: list, scope_name: str, fmt: str) -> tuple[bytes, str, str]:
    d = await gather_aggregate(db, org_ids, scope_name)
    safe = "".join(ch for ch in scope_name if ch.isalnum() or ch in " -_").strip().replace(" ", "_") or "portfolio"
    base = f"{safe}_portfolio_{datetime.now(timezone.utc):%Y%m%d}"
    if fmt == "docx":
        return (render_aggregate_docx(d),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{base}.docx")
    return (render_aggregate_pdf(d), "application/pdf", f"{base}.pdf")


async def build_report(db: AsyncSession, org_id, report_type: str, fmt: str) -> tuple[bytes, str, str]:
    """Returns (content_bytes, mime_type, filename)."""
    d = await gather(db, org_id)
    safe_client = "".join(ch for ch in d["client"] if ch.isalnum() or ch in " -_").strip().replace(" ", "_") or "client"
    base = f"{safe_client}_{report_type}_{datetime.now(timezone.utc):%Y%m%d}"
    if report_type == "evidence":
        return (render_evidence_zip(d), "application/zip", f"{base}_evidence_bundle.zip")
    if fmt == "docx":
        return (render_docx(report_type, d),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{base}.docx")
    return (render_pdf(report_type, d), "application/pdf", f"{base}.pdf")
