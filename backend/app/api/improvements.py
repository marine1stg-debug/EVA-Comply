"""Internal Improvement / Request log (EVA team only).

Super Admin (incl. dev/tester super-admin accounts) can record things to fix or
improve, attach screenshots, browse the log, update status, copy a single
request, or export them all to a Word document (TOC + list + details + images).

Never exposed to client tenants - the router is gated to Super Admin.
"""
import os
import uuid
import tempfile
from datetime import datetime, timezone
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from starlette.background import BackgroundTask

from app.core.database import get_db
from app.core.config import settings
from app.core.upload_guard import validate_upload
from app.api.auth import get_current_user
from app.models.user import User, UserRole
from app.models.improvement import ImprovementRequest, ImprovementAttachment

router = APIRouter()

UPLOAD_DIR = os.path.join(settings.STORAGE_LOCAL_PATH, "improvements")
MAX_BYTES = 12 * 1024 * 1024
CATEGORIES = {"bug", "idea", "question", "other"}
PRIORITIES = {"low", "medium", "high"}
STATUSES = {"open", "in_progress", "done", "wont_fix"}

CATEGORY_LABEL = {"bug": "Bug / fix", "idea": "Improvement", "question": "Question", "other": "Other"}
STATUS_LABEL = {"open": "Open", "in_progress": "In progress", "done": "Implemented", "wont_fix": "Won't fix"}


async def require_super(current_user: User = Depends(get_current_user)) -> User:
    if current_user.role != UserRole.super_admin:
        raise HTTPException(status_code=403, detail="This tool is restricted to Super Admins")
    return current_user


def _ext_for(original: Optional[str], content_type: Optional[str]) -> str:
    if original and "." in original:
        e = original.rsplit(".", 1)[-1].lower()
        if e.isalnum() and len(e) <= 5:
            return e
    sub = (content_type or "image/png").split("/")[-1].lower()
    return {"jpeg": "jpg"}.get(sub, sub) or "png"


def _save_attachment(req_id: uuid.UUID, up: UploadFile, data: bytes) -> ImprovementAttachment:
    validate_upload(up.filename, data, "support")  # images/docs allowed; blocks executables/scripts
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    ext = _ext_for(up.filename, up.content_type)
    stored = f"{uuid.uuid4().hex}.{ext}"
    with open(os.path.join(UPLOAD_DIR, stored), "wb") as fh:
        fh.write(data)
    return ImprovementAttachment(
        request_id=req_id, filename=stored,
        content_type=up.content_type or "image/png",
        original_name=up.filename,
    )


def _serialize(r: ImprovementRequest) -> dict:
    return {
        "id": str(r.id),
        "title": r.title,
        "body": r.body or "",
        "category": r.category,
        "category_label": CATEGORY_LABEL.get(r.category, r.category),
        "priority": r.priority,
        "status": r.status,
        "status_label": STATUS_LABEL.get(r.status, r.status),
        "author_name": r.author_name,
        "author_role": r.author_role,
        "page_url": r.page_url,
        "resolution_note": r.resolution_note or "",
        "created_at": r.created_at.strftime("%b %d, %Y · %H:%M") if r.created_at else None,
        "created_iso": r.created_at.isoformat() if r.created_at else None,
        "attachments": [{"id": str(a.id), "content_type": a.content_type} for a in (r.attachments or [])],
    }


@router.get("/")
async def list_requests(
    status: Optional[str] = None,
    user: User = Depends(require_super),
    db: AsyncSession = Depends(get_db),
):
    q = select(ImprovementRequest).order_by(ImprovementRequest.created_at.desc())
    if status and status in STATUSES:
        q = q.where(ImprovementRequest.status == status)
    rows = (await db.execute(q)).scalars().all()
    return {"requests": [_serialize(r) for r in rows]}


@router.post("/")
async def create_request(
    title: str = Form(...),
    body: str = Form(""),
    category: str = Form("bug"),
    priority: str = Form("medium"),
    page_url: str = Form(""),
    files: List[UploadFile] = File(default=[]),
    user: User = Depends(require_super),
    db: AsyncSession = Depends(get_db),
):
    if not title.strip():
        raise HTTPException(status_code=400, detail="A title is required")
    req = ImprovementRequest(
        author_id=user.id, author_name=user.display_name, author_role=user.role.value,
        title=title.strip()[:300], body=(body.strip() or None),
        category=category if category in CATEGORIES else "bug",
        priority=priority if priority in PRIORITIES else "medium",
        status="open", page_url=(page_url.strip() or None),
    )
    db.add(req)
    await db.flush()  # need req.id for attachments
    for up in files or []:
        data = await up.read()
        if not data:
            continue
        if len(data) > MAX_BYTES:
            raise HTTPException(status_code=413, detail="Attachment exceeds 12 MB")
        db.add(_save_attachment(req.id, up, data))
    await db.commit()
    await db.refresh(req)
    return _serialize(req)


async def _get(db: AsyncSession, request_id: str) -> ImprovementRequest:
    try:
        rid = uuid.UUID(request_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="Request not found")
    r = (await db.execute(
        select(ImprovementRequest).where(ImprovementRequest.id == rid)
    )).scalar_one_or_none()
    if not r:
        raise HTTPException(status_code=404, detail="Request not found")
    return r


# ── Word export (declared before /{request_id} so the literal path wins) ──────
def _build_docx(rows: List[ImprovementRequest]) -> str:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    run = title.add_run("EVA Comply — Improvement & Request Log")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0D, 0x2A, 0x43)
    sub = doc.add_paragraph()
    sr = sub.add_run(f"Exported {datetime.now(timezone.utc).strftime('%b %d, %Y')} · {len(rows)} request(s)")
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x5A, 0x6B, 0x7B)

    # Table of contents field (populates when opened in Word → Update Field).
    doc.add_paragraph()
    h = doc.add_paragraph()
    hr = h.add_run("Contents")
    hr.bold = True
    hr.font.size = Pt(14)
    p = doc.add_paragraph()
    r1 = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = 'TOC \\o "1-1" \\h \\z \\u'
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate")
    tt = OxmlElement("w:t"); tt.text = "Right-click here and choose “Update Field” to build the table of contents."
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    for el in (fb, it, fs, tt, fe):
        r1._r.append(el)

    doc.add_page_break()

    for i, r in enumerate(rows, 1):
        doc.add_heading(f"{i}. {r.title}", level=1)
        meta = (
            f"Status: {STATUS_LABEL.get(r.status, r.status)}   |   "
            f"Type: {CATEGORY_LABEL.get(r.category, r.category)}   |   "
            f"Priority: {r.priority.capitalize()}"
        )
        mp = doc.add_paragraph(); mr = mp.add_run(meta); mr.font.size = Pt(9.5); mr.font.color.rgb = RGBColor(0x5A, 0x6B, 0x7B)
        who = f"By {r.author_name}"
        if r.author_role:
            who += f" ({r.author_role.replace('_', ' ')})"
        if r.created_at:
            who += f" · {r.created_at.strftime('%b %d, %Y %H:%M')}"
        wp = doc.add_paragraph(); wr = wp.add_run(who); wr.font.size = Pt(9.5); wr.font.color.rgb = RGBColor(0x5A, 0x6B, 0x7B)
        if r.page_url:
            pp = doc.add_paragraph(); pr = pp.add_run(f"Where: {r.page_url}"); pr.font.size = Pt(9.5); pr.font.color.rgb = RGBColor(0x5A, 0x6B, 0x7B)

        doc.add_paragraph(r.body or "(no description)")

        if r.resolution_note:
            rp = doc.add_paragraph()
            lbl = rp.add_run("Resolution: ")
            lbl.bold = True
            lbl.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
            rp.add_run(r.resolution_note)

        for a in (r.attachments or []):
            if not (a.content_type or "").startswith("image/"):
                continue
            fp = os.path.join(UPLOAD_DIR, a.filename)
            if not os.path.isfile(fp):
                continue
            try:
                doc.add_picture(fp, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

        if i < len(rows):
            doc.add_paragraph()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.close()
    doc.save(tmp.name)
    return tmp.name


@router.get("/export")
async def export_requests(
    user: User = Depends(require_super),
    db: AsyncSession = Depends(get_db),
):
    rows = (await db.execute(
        select(ImprovementRequest).order_by(ImprovementRequest.created_at.desc())
    )).scalars().all()
    path = _build_docx(rows)
    fname = f"EVA_Improvement_Log_{datetime.now(timezone.utc).strftime('%Y%m%d')}.docx"
    return FileResponse(
        path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=fname,
        background=BackgroundTask(lambda: os.path.exists(path) and os.remove(path)),
    )


@router.get("/{request_id}")
async def get_request(
    request_id: str,
    user: User = Depends(require_super),
    db: AsyncSession = Depends(get_db),
):
    return _serialize(await _get(db, request_id))


class UpdateBody(BaseModel):
    status: Optional[str] = None
    priority: Optional[str] = None
    category: Optional[str] = None
    title: Optional[str] = None
    body: Optional[str] = None
    resolution_note: Optional[str] = None


@router.patch("/{request_id}")
async def update_request(
    request_id: str, b: UpdateBody,
    user: User = Depends(require_super),
    db: AsyncSession = Depends(get_db),
):
    r = await _get(db, request_id)
    if b.status is not None and b.status in STATUSES:
        r.status = b.status
    if b.priority is not None and b.priority in PRIORITIES:
        r.priority = b.priority
    if b.category is not None and b.category in CATEGORIES:
        r.category = b.category
    if b.title is not None and b.title.strip():
        r.title = b.title.strip()[:300]
    if b.body is not None:
        r.body = b.body.strip() or None
    if b.resolution_note is not None:
        r.resolution_note = b.resolution_note.strip() or None
    await db.commit()
    await db.refresh(r)
    return _serialize(r)


@router.post("/{request_id}/attachments")
async def add_attachment(
    request_id: str,
    files: List[UploadFile] = File(...),
    user: User = Depends(require_super),
    db: AsyncSession = Depends(get_db),
):
    r = await _get(db, request_id)
    for up in files:
        data = await up.read()
        if not data:
            continue
        if len(data) > MAX_BYTES:
            raise HTTPException(status_code=413, detail="Attachment exceeds 12 MB")
        db.add(_save_attachment(r.id, up, data))
    await db.commit()
    await db.refresh(r)
    return _serialize(r)


@router.get("/{request_id}/attachments/{att_id}")
async def get_attachment(
    request_id: str, att_id: str,
    user: User = Depends(require_super),
    db: AsyncSession = Depends(get_db),
):
    try:
        aid = uuid.UUID(att_id)
    except ValueError:
        raise HTTPException(status_code=404, detail="Attachment not found")
    a = (await db.execute(
        select(ImprovementAttachment).where(ImprovementAttachment.id == aid)
    )).scalar_one_or_none()
    if not a or str(a.request_id) != str((await _get(db, request_id)).id):
        raise HTTPException(status_code=404, detail="Attachment not found")
    fp = os.path.join(UPLOAD_DIR, a.filename)
    if not os.path.isfile(fp):
        raise HTTPException(status_code=404, detail="File missing")
    return FileResponse(fp, media_type=a.content_type or "image/png", filename=a.original_name or a.filename)


@router.delete("/{request_id}")
async def delete_request(
    request_id: str,
    user: User = Depends(require_super),
    db: AsyncSession = Depends(get_db),
):
    r = await _get(db, request_id)
    for a in (r.attachments or []):
        fp = os.path.join(UPLOAD_DIR, a.filename)
        try:
            if os.path.isfile(fp):
                os.remove(fp)
        except OSError:
            pass
    await db.delete(r)
    await db.commit()
    return {"ok": True}
