"""Generate French (.fr.docx) versions of every policy template, applying a
shared Québécois translation memory. Re-runnable; preserves structure/tables."""
import glob, os, re
from docx import Document

# Reuse the Access Control translations (the rich template) ...
_src = open(os.path.join(os.path.dirname(__file__), "make_fr_policy_access_control.py")).read()
_ns = {}
exec(_src[_src.index("TRANS = {"):_src.index("\ndef set_text")], _ns)
TRANS = dict(_ns["TRANS"])

# ... plus the shared boilerplate + per-policy titles/purpose for the other 17.
PURPOSE_TAIL = (" Elle appuie les obligations de conformité de [Organisation] et protège "
                "la confidentialité, l'intégrité et la disponibilité de l'information.")
EXTRA = {
 "Policy Template": "Modèle de politique",
 "Document control": "Contrôle du document",
 "1. Purpose": "1. Objet",
 "2. Scope": "2. Portée",
 "3. Policy statements": "3. Énoncés de politique",
 "4. Roles & responsibilities": "4. Rôles et responsabilités",
 "5. Compliance, exceptions & enforcement": "5. Conformité, exceptions et application",
 "6. Review & maintenance": "6. Révision et maintenance",
 "7. Evidence to retain": "7. Preuves à conserver",
 "This policy applies to all employees, contractors, and third parties who access [Organization]'s information, systems, applications and facilities, and to all systems that store, process or transmit organizational information.":
   "La présente politique s'applique à tous les employés, sous-traitants et tiers qui accèdent à l'information, aux systèmes, aux applications et aux installations de [Organisation], ainsi qu'à tous les systèmes qui stockent, traitent ou transmettent de l'information organisationnelle.",
 "[Organization] shall implement and maintain the controls described by this policy.":
   "[Organisation] doit mettre en œuvre et maintenir les contrôles décrits dans la présente politique.",
 "Responsibilities shall be clearly assigned to specific individuals or roles.":
   "Les responsabilités doivent être clairement attribuées à des personnes ou à des rôles précis.",
 "Activities required by this policy shall be performed, documented and retained as evidence.":
   "Les activités exigées par la présente politique doivent être réalisées, documentées et conservées à titre de preuve.",
 "Exceptions shall be risk-assessed, approved by the policy owner, and time-bound.":
   "Les exceptions doivent faire l'objet d'une évaluation des risques, être approuvées par le propriétaire de la politique et être limitées dans le temps.",
 "[Add the specific control statements that apply to your environment here.]":
   "[Ajoutez ici les énoncés de contrôle précis qui s'appliquent à votre environnement.]",
 "Executive sponsor — approves the policy and provides resources.":
   "Parrain de la haute direction — approuve la politique et fournit les ressources.",
 "Policy owner — maintains the policy and oversees its implementation.":
   "Propriétaire de la politique — maintient la politique et supervise sa mise en œuvre.",
 "System owners / managers — implement and operate the controls.":
   "Propriétaires / gestionnaires de systèmes — mettent en œuvre et exploitent les contrôles.",
 "All users — comply with the policy and report issues.":
   "Tous les utilisateurs — se conforment à la politique et signalent les problèmes.",
 "Compliance with this policy is mandatory. Non-compliance may result in disciplinary action. Exceptions must be documented, risk-assessed and approved by the policy owner, and reviewed at least [annually].":
   "Le respect de la présente politique est obligatoire. Tout manquement peut entraîner des mesures disciplinaires. Les exceptions doivent être documentées, faire l'objet d'une évaluation des risques et être approuvées par le propriétaire de la politique, puis révisées au moins [annuellement].",
 "This policy shall be reviewed at least annually, or following significant changes to the business, technology or threat landscape, and updated as required.":
   "La présente politique doit être révisée au moins une fois par année, ou à la suite de changements importants touchant l'entreprise, la technologie ou le paysage des menaces, et mise à jour au besoin.",
 "This signed and dated policy.": "La présente politique, signée et datée.",
 "Records showing the controls are operating (e.g. approvals, reviews, logs, tickets).":
   "Des registres démontrant que les contrôles fonctionnent (p. ex. approbations, revues, journaux, billets).",
 "Exception records and their approvals.": "Les dossiers d'exception et leurs approbations.",
 "Policy owner": "Propriétaire de la politique",
 "[Name / role]": "[Nom / rôle]",
 "Approved by": "Approuvée par",
 "Effective date": "Date d'entrée en vigueur",
 "Last reviewed": "Dernière révision",
 "Review cycle": "Cycle de révision",
 "Annual": "Annuel",
 "Internal": "Interne",
 # Titles
 "Audit & Accountability Policy": "Politique d'audit et de responsabilisation",
 "Configuration Management Policy": "Politique de gestion de la configuration",
 "Identification & Authentication Policy": "Politique d'identification et d'authentification",
 "Incident Response Policy & Plan": "Politique et plan de réponse aux incidents",
 "Information Security Policy": "Politique de sécurité de l'information",
 "Information Security Program Plan (SSP)": "Plan du programme de sécurité de l'information (SSP)",
 "Media Protection Policy": "Politique de protection des supports",
 "Personnel Security Policy": "Politique de sécurité du personnel",
 "Physical Protection Policy": "Politique de protection physique",
 "Risk Assessment Policy": "Politique d'évaluation des risques",
 "Security Assessment & Continuous Monitoring Policy": "Politique d'évaluation de la sécurité et de surveillance continue",
 "Security Awareness & Training Policy": "Politique de sensibilisation et de formation en sécurité",
 "Supply Chain Risk Management Policy": "Politique de gestion des risques liés à la chaîne d'approvisionnement",
 "System & Communications Protection Policy": "Politique de protection des systèmes et des communications",
 "System & Information Integrity Policy": "Politique d'intégrité des systèmes et de l'information",
 "System Maintenance Policy": "Politique de maintenance des systèmes",
 "System & Services Acquisition Policy": "Politique d'acquisition des systèmes et des services",
 # Per-policy purpose sentences (lead clause + shared tail)
 "The purpose of this policy is to define how security-relevant events are logged, protected, reviewed and retained for accountability. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de définir comment les événements pertinents pour la sécurité sont journalisés, protégés, examinés et conservés à des fins de responsabilisation." + PURPOSE_TAIL,
 "The purpose of this policy is to establish secure baseline configurations and control changes to systems over their lifecycle. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet d'établir des configurations de référence sécurisées et de contrôler les changements aux systèmes tout au long de leur cycle de vie." + PURPOSE_TAIL,
 "The purpose of this policy is to uniquely identify users and devices and authenticate them before granting access. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet d'identifier de façon unique les utilisateurs et les appareils et de les authentifier avant d'accorder l'accès." + PURPOSE_TAIL,
 "The purpose of this policy is to prepare for, detect, respond to, and recover from cybersecurity incidents. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de se préparer aux incidents de cybersécurité, de les détecter, d'y répondre et de s'en rétablir." + PURPOSE_TAIL,
 "The purpose of this policy is to set the organization's overall expectations for protecting the confidentiality, integrity and availability of information. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet d'établir les attentes générales de l'organisation en matière de protection de la confidentialité, de l'intégrité et de la disponibilité de l'information." + PURPOSE_TAIL,
 "The purpose of this policy is to describe the organization's information security program, system boundary, and how controls are implemented. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de décrire le programme de sécurité de l'information de l'organisation, le périmètre du système et la manière dont les contrôles sont mis en œuvre." + PURPOSE_TAIL,
 "The purpose of this policy is to protect, handle, transport and sanitize digital and physical media containing sensitive information. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de protéger, de manipuler, de transporter et d'assainir les supports numériques et physiques contenant de l'information sensible." + PURPOSE_TAIL,
 "The purpose of this policy is to manage security risk associated with personnel through screening, agreements, transfers and terminations. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de gérer le risque de sécurité associé au personnel au moyen de la vérification, des ententes, des mutations et des départs." + PURPOSE_TAIL,
 "The purpose of this policy is to control physical access to facilities, equipment and supporting infrastructure. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de contrôler l'accès physique aux installations, à l'équipement et à l'infrastructure de soutien." + PURPOSE_TAIL,
 "The purpose of this policy is to identify, assess and prioritize risks to operations, assets and individuals on a recurring basis. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de cerner, d'évaluer et de prioriser, de façon récurrente, les risques pour les opérations, les actifs et les personnes." + PURPOSE_TAIL,
 "The purpose of this policy is to assess controls, manage findings and continuously monitor the security posture. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet d'évaluer les contrôles, de gérer les constats et de surveiller en continu la posture de sécurité." + PURPOSE_TAIL,
 "The purpose of this policy is to ensure personnel receive security awareness and role-based training appropriate to their duties. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de veiller à ce que le personnel reçoive une sensibilisation à la sécurité et une formation adaptée à ses fonctions." + PURPOSE_TAIL,
 "The purpose of this policy is to manage cybersecurity risk arising from suppliers, products and the supply chain. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de gérer le risque de cybersécurité découlant des fournisseurs, des produits et de la chaîne d'approvisionnement." + PURPOSE_TAIL,
 "The purpose of this policy is to protect information in transit and at rest and segregate and defend the network. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de protéger l'information en transit et au repos, et de segmenter et défendre le réseau." + PURPOSE_TAIL,
 "The purpose of this policy is to identify, report and correct flaws, and protect systems from malicious code and threats. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet de cerner, de signaler et de corriger les failles, et de protéger les systèmes contre le code malveillant et les menaces." + PURPOSE_TAIL,
 "The purpose of this policy is to perform system maintenance securely and control maintenance tools, personnel and remote sessions. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet d'effectuer la maintenance des systèmes de façon sécurisée et de contrôler les outils, le personnel et les sessions à distance de maintenance." + PURPOSE_TAIL,
 "The purpose of this policy is to build security into the acquisition and development of systems, components and services. It supports [Organization]'s compliance obligations and protects the confidentiality, integrity and availability of information.":
   "La présente politique a pour objet d'intégrer la sécurité à l'acquisition et au développement des systèmes, des composants et des services." + PURPOSE_TAIL,
}
TRANS.update(EXTRA)

def set_text(p, txt):
    if not p.runs:
        p.add_run(txt); return
    p.runs[0].text = txt
    for r in p.runs[1:]:
        r.text = ""

def trivial(s):
    return bool(re.fullmatch(r"[\d\.\sA-Z()\-/&,]+|\[1\.0\]|\[Date\]", s)) and not any(c.islower() for c in s)

def do(paras, miss):
    for p in paras:
        t = p.text.strip()
        if not t: continue
        if t in TRANS and TRANS[t]:
            set_text(p, TRANS[t])
        elif not trivial(t):
            miss.append(t)

docs = [d for d in sorted(glob.glob("policy_library/*.docx")) if not d.endswith(".fr.docx")]
grand = 0
for path in docs:
    d = Document(path); miss = []
    do(d.paragraphs, miss)
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                do(cell.paragraphs, miss)
    out = path[:-5] + ".fr.docx"
    d.save(out)
    grand += len(miss)
    flag = "" if not miss else f"  !! {len(miss)} untranslated"
    print(f"  {os.path.basename(out)}{flag}")
    for m in miss[:6]: print("       -", repr(m[:80]))
print(f"\nTotal untranslated (non-trivial) across all docs: {grand}")
