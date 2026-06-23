"""Build a policy .docx matching the Access Control Policy's format:
Heading 2 (Arial bold, blue 2E5FA3), bordered tables with navy 1B3A6B header
shading + white header text, metadata/roles/mapping/revision tables."""
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "policy_library/Access_Control_Policy.docx"
NAVY = "1B3A6B"
BLUE = RGBColor(0x2E, 0x5F, 0xA3)

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _borders(table):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'auto')
        b.append(e)
    tblPr.append(b)

def _cell(cell, text, bold=False, white=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text); r.bold = bold
    r.font.name = "Arial"; r.font.size = Pt(size)
    if white: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

def new_doc():
    d = Document(BASE)
    # wipe body (paragraphs + tables) but keep styles/theme
    body = d.element.body
    for child in list(body):
        if child.tag in (qn('w:p'), qn('w:tbl')):
            body.remove(child)
    return d

def heading(d, text):
    p = d.add_paragraph()
    pf = p.paragraph_format; pf.space_before=Pt(10); pf.space_after=Pt(4)
    r = p.add_run(text); r.bold=True; r.font.name="Arial"; r.font.size=Pt(13); r.font.color.rgb=BLUE
    return p

def body(d, text, size=10.5):
    p = d.add_paragraph()
    r = p.add_run(text); r.font.name="Arial"; r.font.size=Pt(size)
    return p

def bullet(d, text, size=10.5):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    r = p.add_run("•  " + text); r.font.name="Arial"; r.font.size=Pt(size)
    return p

def ps(d, label, text):
    p = d.add_paragraph()
    r = p.add_run(label); r.bold=True; r.font.name="Arial"; r.font.size=Pt(10.5); r.font.color.rgb=BLUE
    body(d, text)

def table(d, rows, header=True, widths=None):
    cols = len(rows[0])
    t = d.add_table(rows=len(rows), cols=cols)
    _borders(t)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            ish = header and ri == 0
            _cell(t.rows[ri].cells[ci], str(val), bold=ish, white=ish)
            if ish: _shade(t.rows[ri].cells[ci], NAVY)
    d.add_paragraph()
    return t

def title_block(d, org_ph, policy_title):
    t = d.add_table(rows=1, cols=2); _borders(t)
    _cell(t.rows[0].cells[0], "[INSERT ORGANIZATION LOGO]", size=10)
    c2 = t.rows[0].cells[1]; c2.text=""
    p1=c2.paragraphs[0]; r1=p1.add_run(org_ph); r1.bold=True; r1.font.name="Arial"; r1.font.size=Pt(12)
    p2=c2.add_paragraph(); r2=p2.add_run(policy_title); r2.bold=True; r2.font.name="Arial"; r2.font.size=Pt(14); r2.font.color.rgb=BLUE
    d.add_paragraph()

DEFAULT_VIOLATIONS = [
    "Remediation of the deficiency within a timeframe set by the policy owner based on risk.",
    "Disciplinary action up to and including termination of employment or contract, consistent with [Organization Name]'s disciplinary procedures.",
    "Referral to law enforcement and notification to government contracting officers where CUI or criminal activity is involved, as required by the applicable contract or regulation.",
]
DEFAULT_EXCEPTIONS = [
    "The requester submits a written exception to the policy owner describing the requirement that cannot be met, the business justification, the risk introduced, and proposed compensating controls.",
    "The policy owner assesses the risk and prepares a recommendation; [Senior Leadership / CISO] approves or denies the exception in writing.",
    "Approved exceptions are logged in the Exception Register with an expiry of no more than [organization-defined: 12 months; recommended: 12 months], reviewed upon expiry, and reflected as a plan of action and milestones (POA&M) item in the SSP where CUI systems are affected.",
]

def _kv_def(d, term, defn):
    p = d.add_paragraph()
    r = p.add_run(term + ": "); r.bold=True; r.font.name="Arial"; r.font.size=Pt(10.5)
    r2 = p.add_run(defn); r2.font.name="Arial"; r2.font.size=Pt(10.5)

def _evidence(d, text):
    p = d.add_paragraph()
    r = p.add_run("Audit evidence: "); r.bold=True; r.font.name="Arial"; r.font.size=Pt(10.5)
    r2 = p.add_run(text); r2.font.name="Arial"; r2.font.size=Pt(10.5)

def build(spec):
    d = new_doc()
    owner = spec.get("owner", "[IT Manager / ISSO]")
    title_block(d, "[Organization Name]", spec["title"])
    table(d, [
        ["Field","Value","Field","Value"],
        ["Policy Name", spec["title"], "Version", "[1.0]"],
        ["Policy Owner", owner, "Approved By", "[CEO / CISO]"],
        ["Effective Date","[YYYY-MM-DD]","Last Reviewed","[YYYY-MM-DD]"],
        ["Review Cadence","Annual (or upon major change)","Classification","[CUI / Internal]"],
    ])
    heading(d, "1.  Purpose"); body(d, spec["purpose"])
    heading(d, "2.  Scope"); body(d, "This policy applies to:")
    for b in spec["scope"]: bullet(d, b)
    body(d, spec.get("scope_tail", "Where CUI is involved, the requirements in this policy are mandatory and supersede convenience or operational preferences."))
    heading(d, "3.  Definitions")
    for term, defn in spec["definitions"]: _kv_def(d, term, defn)
    heading(d, "4.  Roles & Responsibilities")
    table(d, [["Role","Key Responsibilities"]] + spec["roles"])
    heading(d, "5.  Policy Statements")
    for label, text in spec["statements"]: ps(d, label, text)
    heading(d, "6.  Procedures & Audit Evidence")
    for sub, bullets, ev in spec["procedures"]:
        body(d, sub)
        for b in bullets: bullet(d, b)
        _evidence(d, ev)
    heading(d, "7.  Compliance, Monitoring & Enforcement")
    body(d, f"The {owner} shall review compliance with this policy at least [quarterly] " + spec.get("compliance_focus", "through documented reviews, configuration checks, and audits.") + " Results shall be reported to [Senior Leadership / CISO].")
    body(d, "Violations may result in:")
    for v in spec.get("violations", DEFAULT_VIOLATIONS): bullet(d, v)
    heading(d, "8.  Exceptions")
    body(d, "Exceptions to this policy must be rare. To request one:")
    for b in DEFAULT_EXCEPTIONS: bullet(d, b)
    heading(d, "9.  Related Documents")
    for b in spec["related"]: bullet(d, b)
    heading(d, "10.  Control Mapping")
    body(d, spec["mapping_intro"])
    table(d, [["Control ID","Policy Statement(s)","Expected Audit Evidence"]] + spec["mapping"])
    heading(d, "11.  Revision History")
    table(d, [["Version","Date","Author / Role","Summary of Changes"],
              ["1.0","[YYYY-MM-DD]", owner, "Initial policy release."]])
    d.save(spec["path"])
    return d

# ── French generation ────────────────────────────────────────────────────────
FR_SECTIONS = ["1.  Objet","2.  Portée","3.  Définitions","4.  Rôles et responsabilités",
    "5.  Énoncés de politique","6.  Procédures et preuves d'audit",
    "7.  Conformité, surveillance et application","8.  Exceptions","9.  Documents connexes",
    "10.  Correspondance des contrôles","11.  Historique des révisions"]
FR_VIOLATIONS = [
    "La correction de la lacune dans un délai fixé par le propriétaire de la politique en fonction du risque.",
    "Des mesures disciplinaires pouvant aller jusqu'au congédiement ou à la résiliation du contrat, conformément aux procédures disciplinaires de [Nom de l'organisation].",
    "Un renvoi aux forces de l'ordre et un avis aux agents de contrats gouvernementaux lorsque des CUI ou une activité criminelle sont en cause, comme l'exige le contrat ou le règlement applicable.",
]
FR_EXCEPTIONS = [
    "Le demandeur soumet une demande d'exception écrite au propriétaire de la politique décrivant l'exigence qui ne peut être respectée, la justification d'affaires, le risque introduit et les contrôles compensatoires proposés.",
    "Le propriétaire de la politique évalue le risque et prépare une recommandation; la [haute direction / le RSSI] approuve ou refuse l'exception par écrit.",
    "Les exceptions approuvées sont consignées au registre des exceptions avec une expiration d'au plus [défini par l'organisation : 12 mois; recommandé : 12 mois], révisées à l'échéance et reflétées comme élément d'un plan d'action et de jalons (POA&M) dans le SSP lorsque des systèmes contenant des CUI sont touchés.",
]

def build_fr(spec):
    d = new_doc()
    owner = spec.get("owner", "[Gestionnaire TI / ISSO]")
    title_block(d, "[Nom de l'organisation]", spec["title"])
    table(d, [
        ["Champ","Valeur","Champ","Valeur"],
        ["Nom de la politique", spec["title"], "Version", "[1.0]"],
        ["Propriétaire de la politique", owner, "Approuvée par", "[PDG / RSSI]"],
        ["Date d'entrée en vigueur","[AAAA-MM-JJ]","Dernière révision","[AAAA-MM-JJ]"],
        ["Fréquence de révision","Annuelle (ou lors d'un changement majeur)","Classification","[CUI / Interne]"],
    ])
    S = FR_SECTIONS
    heading(d, S[0]); body(d, spec["purpose"])
    heading(d, S[1]); body(d, "La présente politique s'applique à :")
    for b in spec["scope"]: bullet(d, b)
    body(d, spec.get("scope_tail", "Lorsque des CUI sont en cause, les exigences de la présente politique sont obligatoires et prévalent sur la commodité ou les préférences opérationnelles."))
    heading(d, S[2])
    for term, defn in spec["definitions"]: _kv_def(d, term, defn)
    heading(d, S[3])
    table(d, [["Rôle","Principales responsabilités"]] + spec["roles"])
    heading(d, S[4])
    for label, text in spec["statements"]: ps(d, label, text)
    heading(d, S[5])
    for sub, bullets, ev in spec["procedures"]:
        body(d, sub)
        for b in bullets: bullet(d, b)
        p=d.add_paragraph(); r=p.add_run("Preuves d'audit : "); r.bold=True; r.font.name="Arial"; r.font.size=Pt(10.5)
        r2=p.add_run(ev); r2.font.name="Arial"; r2.font.size=Pt(10.5)
    heading(d, S[6])
    body(d, f"Le {owner} doit examiner la conformité à la présente politique au moins [trimestriellement] " + spec.get("compliance_focus","au moyen de revues documentées, de vérifications de configuration et d'audits.") + " Les résultats doivent être communiqués à la [haute direction / au RSSI].")
    body(d, "Les manquements peuvent entraîner :")
    for v in spec.get("violations", FR_VIOLATIONS): bullet(d, v)
    heading(d, S[7])
    body(d, "Les exceptions à la présente politique doivent demeurer rares. Pour en demander une :")
    for b in FR_EXCEPTIONS: bullet(d, b)
    heading(d, S[8])
    for b in spec["related"]: bullet(d, b)
    heading(d, S[9]); body(d, spec["mapping_intro"])
    table(d, [["ID du contrôle","Énoncé(s) de politique","Preuves d'audit attendues"]] + spec["mapping"])
    heading(d, S[10])
    table(d, [["Version","Date","Auteur / Rôle","Résumé des modifications"],
              ["1.0","[AAAA-MM-JJ]", owner, "Publication initiale de la politique."]])
    d.save(spec["path"])
    return d
