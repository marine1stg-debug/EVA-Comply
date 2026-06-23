"""Generate the French Access Control Policy from the English template,
preserving structure, styles, tables, and placeholders. Québécois French."""
from docx import Document

TRANS = {
 "1.  Purpose": "1.  Objet",
 "This policy establishes the rules governing who may access [Organization Name]'s systems, networks, and Controlled Unclassified Information (CUI), and what those individuals may do once access is granted. Unauthorized or excessive access is one of the leading causes of data breaches and compliance violations. By enforcing strict account management, least privilege, and session controls, this policy reduces the risk that CUI is accessed, modified, or disclosed without authorization — whether by insiders, compromised credentials, or third parties.":
 "La présente politique établit les règles régissant qui peut accéder aux systèmes, aux réseaux et à l'information contrôlée non classifiée (CUI) de [Nom de l'organisation], et ce que ces personnes peuvent faire une fois l'accès accordé. Les accès non autorisés ou excessifs comptent parmi les principales causes d'atteintes à la protection des données et de manquements à la conformité. En imposant une gestion rigoureuse des comptes, le moindre privilège et des contrôles de session, cette politique réduit le risque que des CUI soient consultés, modifiés ou divulgués sans autorisation — que ce soit par des initiés, des identifiants compromis ou des tiers.",
 "2.  Scope": "2.  Portée",
 "This policy applies to:": "La présente politique s'applique à :",
 "All employees, contractors, consultants, vendors, and any other individuals who access organizational systems or data.":
 "Tous les employés, sous-traitants, consultants, fournisseurs et toute autre personne qui accède aux systèmes ou aux données de l'organisation.",
 "All information systems owned, operated, or managed by [Organization Name], including cloud services, on-premises servers, endpoints, mobile devices, and operational technology that process, store, or transmit CUI.":
 "Tous les systèmes d'information détenus, exploités ou gérés par [Nom de l'organisation], y compris les services infonuagiques, les serveurs sur place, les postes de travail, les appareils mobiles et les technologies opérationnelles qui traitent, stockent ou transmettent des CUI.",
 "Third-party systems that connect to organizational networks or handle CUI on the organization's behalf.":
 "Les systèmes de tiers qui se connectent aux réseaux de l'organisation ou qui traitent des CUI pour le compte de l'organisation.",
 "Where CUI is involved, requirements are non-negotiable and supersede convenience or operational preferences.":
 "Lorsque des CUI sont en cause, les exigences sont non négociables et prévalent sur la commodité ou les préférences opérationnelles.",
 "3.  Definitions": "3.  Définitions",
 "Account: A user identity (username + credential) that grants access to a system or application.":
 "Compte : Une identité d'utilisateur (nom d'utilisateur + identifiant) qui donne accès à un système ou à une application.",
 "CUI (Controlled Unclassified Information): Government information that requires safeguarding under law or regulation (e.g., technical data, export-controlled information). See your System Security Plan (SSP) for the specific CUI categories in scope.":
 "CUI (information contrôlée non classifiée) : Information gouvernementale qui doit être protégée en vertu d'une loi ou d'un règlement (p. ex. données techniques, information visée par le contrôle des exportations). Consultez votre plan de sécurité du système (SSP) pour connaître les catégories de CUI visées.",
 "Least Privilege: Giving a user only the minimum access needed to perform their job — nothing more.":
 "Moindre privilège : N'accorder à un utilisateur que l'accès minimal nécessaire à l'exécution de ses fonctions — rien de plus.",
 "Privileged Account: An account with elevated rights (e.g., administrator, root, system account) that can configure systems, manage other accounts, or override controls.":
 "Compte privilégié : Un compte doté de droits élevés (p. ex. administrateur, racine, compte système) pouvant configurer des systèmes, gérer d'autres comptes ou contourner des contrôles.",
 "Separation of Duties (SoD): Splitting critical tasks between two or more people so no single person can commit fraud or make a critical error undetected.":
 "Séparation des tâches (SoD) : Répartir les tâches critiques entre deux personnes ou plus afin qu'aucune personne seule ne puisse commettre une fraude ou une erreur critique sans être détectée.",
 "Multi-Factor Authentication (MFA): Requiring two or more independent proofs of identity (e.g., password + phone code) before granting access.":
 "Authentification multifacteur (AMF) : Exiger deux preuves d'identité indépendantes ou plus (p. ex. mot de passe + code envoyé au téléphone) avant d'accorder l'accès.",
 "Session Lock: Automatically hiding screen content and requiring re-authentication after a period of inactivity.":
 "Verrouillage de session : Masquer automatiquement le contenu de l'écran et exiger une réauthentification après une période d'inactivité.",
 "Remote Access: Connecting to organizational systems from outside the organization's physical premises (e.g., from home, hotel, or a third-party site).":
 "Accès à distance : Se connecter aux systèmes de l'organisation depuis l'extérieur de ses locaux physiques (p. ex. depuis le domicile, un hôtel ou le site d'un tiers).",
 "MDM (Mobile Device Management): Software that enforces security policies on smartphones and tablets.":
 "GAM (gestion des appareils mobiles) : Logiciel qui applique des politiques de sécurité sur les téléphones intelligents et les tablettes.",
 "External System: Any computer system not owned or controlled by [Organization Name], including personal computers and cloud services not authorized in the SSP.":
 "Système externe : Tout système informatique non détenu ni contrôlé par [Nom de l'organisation], y compris les ordinateurs personnels et les services infonuagiques non autorisés dans le SSP.",
 "4.  Roles & Responsibilities": "4.  Rôles et responsabilités",
 "5.  Policy Statements": "5.  Énoncés de politique",
 "PS-1  Account Authorization": "PS-1  Autorisation des comptes",
 "All user and service accounts must be formally authorized by an account owner before being created. Authorization must be documented (e.g., a signed access request form or ticket). A complete inventory of all accounts — including their type, owner, and access level — must be maintained and kept current at all times.":
 "Tous les comptes d'utilisateur et de service doivent être officiellement autorisés par un propriétaire de compte avant d'être créés. L'autorisation doit être documentée (p. ex. un formulaire de demande d'accès signé ou un billet). Un inventaire complet de tous les comptes — y compris leur type, leur propriétaire et leur niveau d'accès — doit être tenu et maintenu à jour en tout temps.",
 "PS-2  Account Types": "PS-2  Types de comptes",
 "The organization must define and enforce distinct account types (individual user, service/shared, privileged/administrator, guest). Shared accounts are prohibited on systems that process CUI unless technically unavoidable and explicitly approved with compensating controls documented in the SSP.":
 "L'organisation doit définir et appliquer des types de comptes distincts (utilisateur individuel, service/partagé, privilégié/administrateur, invité). Les comptes partagés sont interdits sur les systèmes qui traitent des CUI, sauf s'ils sont techniquement inévitables et explicitement approuvés, avec des contrôles compensatoires documentés dans le SSP.",
 "PS-3  Least Privilege": "PS-3  Moindre privilège",
 "Each account must be provisioned with only the minimum permissions required to perform the user's assigned job function. Access rights must be reviewed [organization-defined: at least quarterly; recommended: quarterly], and any excess privileges removed within [organization-defined: 5 business days; recommended: 5 business days] of discovery.":
 "Chaque compte ne doit être doté que des permissions minimales requises pour exécuter la fonction assignée à l'utilisateur. Les droits d'accès doivent être révisés [défini par l'organisation : au moins trimestriellement; recommandé : trimestriellement], et tout privilège excédentaire doit être retiré dans les [défini par l'organisation : 5 jours ouvrables; recommandé : 5 jours ouvrables] suivant sa découverte.",
 "PS-4  Separation of Duties": "PS-4  Séparation des tâches",
 "Critical or high-risk functions — such as financial transactions above a defined threshold, code promotion to production, and account provisioning — must be divided so that no single individual can complete the action alone. The organization must document which functions require SoD and maintain an up-to-date role matrix.":
 "Les fonctions critiques ou à risque élevé — comme les transactions financières au-delà d'un seuil défini, la mise en production de code et l'approvisionnement des comptes — doivent être réparties de sorte qu'aucune personne seule ne puisse mener l'action à terme. L'organisation doit documenter les fonctions exigeant une séparation des tâches et tenir à jour une matrice des rôles.",
 "PS-5  Privileged Account Controls": "PS-5  Contrôles des comptes privilégiés",
 "Privileged accounts must be separate from standard user accounts. Privileged access must not be used for routine activities (e.g., reading email, browsing the web). Use of privileged accounts must be logged. The number of privileged accounts must be limited to the minimum necessary.":
 "Les comptes privilégiés doivent être distincts des comptes d'utilisateur standard. Les accès privilégiés ne doivent pas servir aux activités courantes (p. ex. lire le courriel, naviguer sur le Web). L'utilisation des comptes privilégiés doit être journalisée. Le nombre de comptes privilégiés doit être limité au minimum nécessaire.",
 "PS-6  Access Termination and Transfer": "PS-6  Cessation et transfert des accès",
 "When an individual leaves the organization or changes roles, their access must be reviewed and adjusted within [organization-defined: 1 business day for termination, 5 business days for role change; recommended: same-day for termination]. HR must notify IT of personnel changes within 1 business day.":
 "Lorsqu'une personne quitte l'organisation ou change de rôle, ses accès doivent être révisés et ajustés dans les [défini par l'organisation : 1 jour ouvrable pour un départ, 5 jours ouvrables pour un changement de rôle; recommandé : le jour même pour un départ]. Les RH doivent aviser les TI des changements de personnel dans un délai de 1 jour ouvrable.",
 "PS-7  Remote Access Controls": "PS-7  Contrôles de l'accès à distance",
 "Remote access to organizational systems must be authorized, encrypted using [organization-defined: TLS 1.2 or higher; VPN with AES-256; recommended: TLS 1.3 or VPN with AES-256], and protected by MFA. Remote access sessions must be monitored. Users must not route organizational traffic through unapproved networks or devices. Remote-access permissions must be reviewed [organization-defined: annually; recommended: annually].":
 "L'accès à distance aux systèmes de l'organisation doit être autorisé, chiffré au moyen de [défini par l'organisation : TLS 1.2 ou plus récent; RPV avec AES-256; recommandé : TLS 1.3 ou RPV avec AES-256] et protégé par l'AMF. Les sessions d'accès à distance doivent être surveillées. Les utilisateurs ne doivent pas acheminer le trafic de l'organisation par des réseaux ou des appareils non approuvés. Les permissions d'accès à distance doivent être révisées [défini par l'organisation : annuellement; recommandé : annuellement].",
 "PS-8  Remote Access via External Networks": "PS-8  Accès à distance par des réseaux externes",
 "Users connecting from external or public networks (e.g., hotel Wi-Fi, coffee shops) must use the organization's approved VPN or secure remote-access gateway for all traffic to organizational systems. Split-tunneling configurations that allow simultaneous access to the internet and organizational systems are prohibited on systems processing CUI unless explicitly approved.":
 "Les utilisateurs qui se connectent depuis des réseaux externes ou publics (p. ex. le Wi-Fi d'un hôtel, un café) doivent utiliser le RPV approuvé de l'organisation ou une passerelle d'accès à distance sécurisée pour tout le trafic vers les systèmes de l'organisation. Les configurations à tunnel partagé qui permettent un accès simultané à Internet et aux systèmes de l'organisation sont interdites sur les systèmes traitant des CUI, sauf approbation explicite.",
 "PS-9  Session Lock": "PS-9  Verrouillage de session",
 "Systems must be configured to automatically lock the screen after no more than [organization-defined: 15 minutes; recommended: 15 minutes] of inactivity. Re-authentication must be required to unlock the session. Users must manually lock their sessions whenever they leave their workstation unattended.":
 "Les systèmes doivent être configurés pour verrouiller automatiquement l'écran après au plus [défini par l'organisation : 15 minutes; recommandé : 15 minutes] d'inactivité. Une réauthentification doit être exigée pour déverrouiller la session. Les utilisateurs doivent verrouiller manuellement leur session chaque fois qu'ils laissent leur poste de travail sans surveillance.",
 "PS-10  Session Termination": "PS-10  Fin de session",
 "Systems must automatically terminate idle sessions after [organization-defined: 30 minutes for high-privilege sessions, 60 minutes for standard sessions; recommended: 30 / 60 minutes]. After termination, the user must fully re-authenticate to resume work.":
 "Les systèmes doivent mettre fin automatiquement aux sessions inactives après [défini par l'organisation : 30 minutes pour les sessions à privilèges élevés, 60 minutes pour les sessions standard; recommandé : 30 / 60 minutes]. Après la fin de la session, l'utilisateur doit se réauthentifier entièrement pour reprendre son travail.",
 "PS-11  Wireless Access": "PS-11  Accès sans fil",
 "Wireless network connections must be authorized before use. Wireless networks used to transmit CUI must use [organization-defined: WPA3-Enterprise or WPA2-Enterprise with AES-CCMP; recommended: WPA3-Enterprise]. Guest wireless networks must be logically separated from networks that carry CUI. Rogue wireless access points must be detected and reported immediately.":
 "Les connexions aux réseaux sans fil doivent être autorisées avant utilisation. Les réseaux sans fil servant à transmettre des CUI doivent utiliser [défini par l'organisation : WPA3-Entreprise ou WPA2-Entreprise avec AES-CCMP; recommandé : WPA3-Entreprise]. Les réseaux sans fil pour invités doivent être logiquement séparés des réseaux qui transportent des CUI. Les points d'accès sans fil non autorisés doivent être détectés et signalés immédiatement.",
 "PS-12  Mobile Device Use": "PS-12  Utilisation des appareils mobiles",
 "Mobile devices (smartphones, tablets, laptops) used to access organizational systems or CUI must be enrolled in the organization's [organization-defined: MDM solution; recommended: Microsoft Intune or equivalent] before use. MDM must enforce encryption, screen lock, remote wipe, and approved-application policies.":
 "Les appareils mobiles (téléphones intelligents, tablettes, ordinateurs portables) servant à accéder aux systèmes de l'organisation ou aux CUI doivent être inscrits à la [défini par l'organisation : solution GAM; recommandé : Microsoft Intune ou l'équivalent] de l'organisation avant utilisation. La GAM doit imposer le chiffrement, le verrouillage d'écran, l'effacement à distance et les politiques d'applications approuvées.",
 "PS-13  Mobile Device Restrictions": "PS-13  Restrictions relatives aux appareils mobiles",
 "CUI must not be stored on personal mobile devices unless those devices are enrolled in MDM and meet all security requirements. Jailbroken or rooted devices must never be used to access organizational systems or CUI. Users must report lost or stolen devices within 4 hours of discovery so that a remote wipe can be initiated.":
 "Les CUI ne doivent pas être stockés sur des appareils mobiles personnels, à moins que ces appareils soient inscrits à la GAM et respectent toutes les exigences de sécurité. Les appareils débridés (jailbreakés) ou rootés ne doivent jamais servir à accéder aux systèmes de l'organisation ni aux CUI. Les utilisateurs doivent signaler tout appareil perdu ou volé dans les 4 heures suivant la découverte afin qu'un effacement à distance puisse être déclenché.",
 "PS-14  External Systems": "PS-14  Systèmes externes",
 "Users must not access, process, store, or transmit CUI on personally owned computers or any system not explicitly authorized in the SSP without a documented approval from the policy owner. Before any external system is permitted to connect to organizational systems, a use agreement and security review must be completed and retained on file.":
 "Les utilisateurs ne doivent pas consulter, traiter, stocker ni transmettre des CUI sur des ordinateurs personnels ou tout système non explicitement autorisé dans le SSP sans une approbation documentée du propriétaire de la politique. Avant qu'un système externe soit autorisé à se connecter aux systèmes de l'organisation, une entente d'utilisation et une revue de sécurité doivent être réalisées et conservées au dossier.",
 "6.  Procedures & Audit Evidence": "6.  Procédures et preuves d'audit",
 "6.1  Account Lifecycle": "6.1  Cycle de vie des comptes",
 "Steps:": "Étapes :",
 "New hire: HR submits access request to IT within 1 business day of start date.":
 "Nouvelle embauche : Les RH soumettent une demande d'accès aux TI dans un délai de 1 jour ouvrable suivant la date d'entrée en fonction.",
 "IT verifies manager approval, creates account with least-privilege permissions, and documents in the account inventory.":
 "Les TI vérifient l'approbation du gestionnaire, créent le compte avec des permissions de moindre privilège et le consignent dans l'inventaire des comptes.",
 "User completes security awareness training before CUI access is granted.":
 "L'utilisateur suit la formation de sensibilisation à la sécurité avant que l'accès aux CUI soit accordé.",
 "Role change: Manager submits updated access request; IT adjusts rights within [5 business days].":
 "Changement de rôle : Le gestionnaire soumet une demande d'accès mise à jour; les TI ajustent les droits dans les [5 jours ouvrables].",
 "Termination: HR notifies IT same day; IT disables account, revokes remote access, and removes CUI access within [1 business day].":
 "Départ : Les RH avisent les TI le jour même; les TI désactivent le compte, révoquent l'accès à distance et retirent l'accès aux CUI dans les [1 jour ouvrable].",
 "Audit evidence: Completed access request forms or tickets; account inventory (dated); termination checklists; training completion records.":
 "Preuves d'audit : Formulaires de demande d'accès ou billets remplis; inventaire des comptes (daté); listes de vérification de départ; relevés d'achèvement de formation.",
 "6.2  Privileged Account Management": "6.2  Gestion des comptes privilégiés",
 "Privileged accounts must be listed and reviewed [quarterly].":
 "Les comptes privilégiés doivent être répertoriés et révisés [trimestriellement].",
 "Admin actions must be logged to a tamper-resistant log store.":
 "Les actions d'administration doivent être journalisées dans un dépôt de journaux résistant à l'altération.",
 "Shared privileged credentials must use a password vault; passwords must be rotated after each use.":
 "Les identifiants privilégiés partagés doivent utiliser un coffre-fort de mots de passe; les mots de passe doivent être changés après chaque utilisation.",
 "Audit evidence: Privileged account list; quarterly review sign-offs; password vault access logs.":
 "Preuves d'audit : Liste des comptes privilégiés; approbations des revues trimestrielles; journaux d'accès au coffre-fort de mots de passe.",
 "6.3  Access Reviews": "6.3  Revues des accès",
 "The IT Manager conducts a formal access review [quarterly] for all CUI systems.":
 "Le gestionnaire TI effectue une revue formelle des accès [trimestriellement] pour tous les systèmes contenant des CUI.",
 "Reviewers certify that each account's access is still needed and appropriately scoped.":
 "Les réviseurs attestent que l'accès de chaque compte est toujours nécessaire et correctement délimité.",
 "Excess access is removed within [5 business days] of review completion.":
 "Les accès excédentaires sont retirés dans les [5 jours ouvrables] suivant la fin de la revue.",
 "Audit evidence: Dated access review reports signed by the IT Manager; evidence of remediation actions.":
 "Preuves d'audit : Rapports de revue des accès datés et signés par le gestionnaire TI; preuves des mesures correctives.",
 "6.4  Remote Access & Session Controls": "6.4  Accès à distance et contrôles de session",
 "VPN and remote-access gateway configurations must be reviewed annually for compliance with PS-7.":
 "Les configurations du RPV et de la passerelle d'accès à distance doivent être révisées annuellement pour assurer la conformité avec PS-7.",
 "Session lock and termination settings must be enforced via Group Policy (GPO) or MDM and verified during [quarterly] configuration audits.":
 "Les paramètres de verrouillage et de fin de session doivent être appliqués au moyen d'une stratégie de groupe (GPO) ou de la GAM et vérifiés lors des audits de configuration [trimestriels].",
 "MFA enrollment must be verified for all remote-access users.":
 "L'inscription à l'AMF doit être vérifiée pour tous les utilisateurs de l'accès à distance.",
 "Audit evidence: VPN configuration exports; GPO/MDM screenshots; MFA enrollment reports; configuration audit logs.":
 "Preuves d'audit : Exportations de la configuration du RPV; captures d'écran GPO/GAM; rapports d'inscription à l'AMF; journaux d'audit de configuration.",
 "6.5  Wireless & Mobile": "6.5  Sans-fil et mobile",
 "Wireless network configurations must be reviewed [annually] or after any network change.":
 "Les configurations des réseaux sans fil doivent être révisées [annuellement] ou après tout changement au réseau.",
 "MDM enrollment status must be verified monthly; non-enrolled devices must be blocked.":
 "L'état d'inscription à la GAM doit être vérifié mensuellement; les appareils non inscrits doivent être bloqués.",
 "Lost-device reports must be logged, and remote wipe must be initiated and confirmed.":
 "Les signalements d'appareils perdus doivent être consignés, et l'effacement à distance doit être déclenché et confirmé.",
 "Audit evidence: Wireless configuration records; MDM enrollment reports; remote-wipe confirmation records.":
 "Preuves d'audit : Registres de configuration sans fil; rapports d'inscription à la GAM; confirmations d'effacement à distance.",
 "6.6  External System Use": "6.6  Utilisation de systèmes externes",
 "Any request to use an external system for CUI must be submitted to the IT Manager.":
 "Toute demande d'utilisation d'un système externe pour des CUI doit être soumise au gestionnaire TI.",
 "The IT Manager completes a security review and issues a written authorization or denial.":
 "Le gestionnaire TI réalise une revue de sécurité et délivre une autorisation ou un refus par écrit.",
 "Users sign an external-system use agreement before access is granted.":
 "Les utilisateurs signent une entente d'utilisation de système externe avant que l'accès soit accordé.",
 "Audit evidence: External-system request forms; security review records; signed use agreements.":
 "Preuves d'audit : Formulaires de demande de système externe; dossiers de revue de sécurité; ententes d'utilisation signées.",
 "7.  Compliance, Monitoring & Enforcement": "7.  Conformité, surveillance et application",
 "The IT Manager shall review compliance with this policy at least [quarterly] through log reviews, access-review reports, and configuration audits. Results shall be reported to [Senior Leadership / CISO].":
 "Le gestionnaire TI doit examiner la conformité à la présente politique au moins [trimestriellement] au moyen de revues de journaux, de rapports de revue des accès et d'audits de configuration. Les résultats doivent être communiqués à la [haute direction / au RSSI].",
 "Violations may result in:": "Les manquements peuvent entraîner :",
 "Immediate suspension of the offending account pending investigation.":
 "La suspension immédiate du compte en cause en attendant une enquête.",
 "Disciplinary action up to and including termination of employment or contract, consistent with [Organization Name]'s disciplinary procedures.":
 "Des mesures disciplinaires pouvant aller jusqu'au congédiement ou à la résiliation du contrat, conformément aux procédures disciplinaires de [Nom de l'organisation].",
 "Referral to law enforcement where criminal activity is suspected.":
 "Un renvoi aux forces de l'ordre lorsqu'une activité criminelle est soupçonnée.",
 "Notification to government contracting officers where CUI is involved, as required by the applicable contract or regulation.":
 "Un avis aux agents de contrats gouvernementaux lorsque des CUI sont en cause, comme l'exige le contrat ou le règlement applicable.",
 "Monitoring activities include: automated account-change alerts, privileged-account login auditing, failed-authentication threshold alerts, and periodic access certification campaigns.":
 "Les activités de surveillance comprennent : les alertes automatisées de modification de compte, l'audit des ouvertures de session des comptes privilégiés, les alertes de seuil d'échecs d'authentification et les campagnes périodiques de certification des accès.",
 "8.  Exceptions": "8.  Exceptions",
 "Exceptions to this policy are discouraged and must be rare. To request an exception:":
 "Les exceptions à la présente politique sont déconseillées et doivent demeurer rares. Pour demander une exception :",
 "The requesting individual or manager submits a written exception request to the IT Manager, describing the specific requirement that cannot be met, the business justification, the risk introduced, and proposed compensating controls.":
 "La personne ou le gestionnaire demandeur soumet une demande d'exception écrite au gestionnaire TI, décrivant l'exigence précise qui ne peut être respectée, la justification d'affaires, le risque introduit et les contrôles compensatoires proposés.",
 "The IT Manager assesses the risk and prepares a recommendation.":
 "Le gestionnaire TI évalue le risque et prépare une recommandation.",
 "[Senior Leadership / CISO] approves or denies the exception in writing.":
 "La [haute direction / le RSSI] approuve ou refuse l'exception par écrit.",
 "Approved exceptions are logged in the Exception Register, include an expiry date of no more than [organization-defined: 12 months; recommended: 12 months], and must be renewed or closed upon expiry.":
 "Les exceptions approuvées sont consignées au registre des exceptions, comportent une date d'expiration d'au plus [défini par l'organisation : 12 mois; recommandé : 12 mois] et doivent être renouvelées ou fermées à l'échéance.",
 "Exceptions affecting CUI systems must be reflected as a plan of action and milestones (POA&M) item in the SSP.":
 "Les exceptions touchant des systèmes contenant des CUI doivent figurer comme élément d'un plan d'action et de jalons (POA&M) dans le SSP.",
 "9.  Related Documents": "9.  Documents connexes",
 "System Security Plan (SSP) — authoritative system inventory, CUI categories, and control implementation details.":
 "Plan de sécurité du système (SSP) — inventaire de référence des systèmes, catégories de CUI et détails de mise en œuvre des contrôles.",
 "Incident Response Policy — procedures when unauthorized access is detected.":
 "Politique de réponse aux incidents — procédures lorsqu'un accès non autorisé est détecté.",
 "Configuration Management Policy — baseline hardening standards that complement session-control requirements.":
 "Politique de gestion de la configuration — normes de durcissement de référence qui complètent les exigences de contrôle de session.",
 "Identification and Authentication Policy — password, MFA, and credential-management requirements.":
 "Politique d'identification et d'authentification — exigences relatives aux mots de passe, à l'AMF et à la gestion des identifiants.",
 "Personnel Security Policy — background screening and onboarding/offboarding procedures.":
 "Politique de sécurité du personnel — procédures de vérification des antécédents et d'entrée/de départ en fonction.",
 "Physical and Environmental Protection Policy — physical access controls that complement logical access controls.":
 "Politique de protection physique et environnementale — contrôles d'accès physique qui complètent les contrôles d'accès logique.",
 "Audit and Accountability Policy — logging, monitoring, and log-retention requirements.":
 "Politique d'audit et de responsabilisation — exigences de journalisation, de surveillance et de conservation des journaux.",
 "10.  Control Mapping": "10.  Correspondance des contrôles",
 "The table below maps each required control to the policy statement(s) that satisfy it and the evidence an auditor will seek. All control IDs from NIST SP 800-171 Rev 3 are included; they are equally applicable to CMMC 2.0 Level 2 and ITSP.10.171 (CPCSC).":
 "Le tableau ci-dessous fait correspondre chaque contrôle requis aux énoncés de politique qui le satisfont ainsi qu'aux preuves qu'un auditeur recherchera. Tous les identifiants de contrôle du NIST SP 800-171 Rev 3 sont inclus; ils s'appliquent également au CMMC 2.0 niveau 2 et à l'ITSP.10.171 (CPCSC).",
 "11.  Revision History": "11.  Historique des révisions",
 # ── Tables ──
 "[Organization Name]": "[Nom de l'organisation]",
 "Access Control Policy": "Politique de contrôle d'accès",
 "Field": "Champ", "Value": "Valeur",
 "Policy Name": "Nom de la politique", "Version": "Version",
 "Policy Owner": "Propriétaire de la politique", "Approved By": "Approuvée par",
 "[IT Manager / ISSO]": "[Gestionnaire TI / ISSO]", "[CEO / CISO]": "[PDG / RSSI]",
 "Effective Date": "Date d'entrée en vigueur", "Last Reviewed": "Dernière révision",
 "[YYYY-MM-DD]": "[AAAA-MM-JJ]",
 "Review Cadence": "Fréquence de révision", "Classification": "Classification",
 "Annual (or upon major change)": "Annuelle (ou lors d'un changement majeur)",
 "[CUI / Internal]": "[CUI / Interne]",
 "Role": "Rôle", "Key Responsibilities": "Principales responsabilités",
 "[IT Manager / ISSO]\nOwn this policy; approve accounts and privileges; review access quarterly; maintain the account inventory; report violations.": None,
 "Own this policy; approve accounts and privileges; review access quarterly; maintain the account inventory; report violations.":
 "Être propriétaire de la présente politique; approuver les comptes et les privilèges; réviser les accès chaque trimestre; tenir l'inventaire des comptes; signaler les manquements.",
 "[System Administrator]": "[Administrateur système]",
 "Implement access controls in systems; execute provisioning and de-provisioning; generate and retain access logs.":
 "Mettre en œuvre les contrôles d'accès dans les systèmes; exécuter l'approvisionnement et le retrait des accès; produire et conserver les journaux d'accès.",
 "[Human Resources]": "[Ressources humaines]",
 "Notify IT within [1 business day] of new hires, role changes, and terminations to trigger account actions.":
 "Aviser les TI dans les [1 jour ouvrable] des nouvelles embauches, des changements de rôle et des départs afin de déclencher les actions sur les comptes.",
 "[Department Manager / Supervisor]": "[Gestionnaire de service / Superviseur]",
 "Authorize access requests for their team; confirm least-privilege assignments; approve exceptions within their scope.":
 "Autoriser les demandes d'accès de leur équipe; confirmer les attributions de moindre privilège; approuver les exceptions relevant de leur portée.",
 "All Users": "Tous les utilisateurs",
 "Use only the accounts and privileges assigned to them; report suspected unauthorized access immediately.":
 "N'utiliser que les comptes et les privilèges qui leur sont attribués; signaler immédiatement tout accès non autorisé soupçonné.",
 "[Senior Leadership / CISO]": "[Haute direction / RSSI]",
 "Approve policy; review compliance metrics; accept residual risk for documented exceptions.":
 "Approuver la politique; examiner les indicateurs de conformité; accepter le risque résiduel pour les exceptions documentées.",
 "Control ID": "ID du contrôle", "Policy Statement(s)": "Énoncé(s) de politique", "Expected Audit Evidence": "Preuves d'audit attendues",
 "PS-1, PS-2, PS-3 — Authorized access; least privilege; account types": "PS-1, PS-2, PS-3 — Accès autorisé; moindre privilège; types de comptes",
 "Account inventory; access request forms; provisioning records": "Inventaire des comptes; formulaires de demande d'accès; registres d'approvisionnement",
 "PS-4, PS-5 — Separation of duties; dual-authorization for privileged actions": "PS-4, PS-5 — Séparation des tâches; double autorisation pour les actions privilégiées",
 "Role matrix showing SoD; privileged-action logs": "Matrice des rôles illustrant la séparation des tâches; journaux des actions privilégiées",
 "PS-3, PS-6 — Least privilege; privileged account restrictions": "PS-3, PS-6 — Moindre privilège; restrictions des comptes privilégiés",
 "Privileged account list; access review reports": "Liste des comptes privilégiés; rapports de revue des accès",
 "PS-7, PS-8 — Remote access controls; encrypted sessions; MFA": "PS-7, PS-8 — Contrôles de l'accès à distance; sessions chiffrées; AMF",
 "VPN/remote-access config; MFA enrollment logs; session-timeout settings": "Configuration du RPV/de l'accès à distance; journaux d'inscription à l'AMF; paramètres d'expiration de session",
 "PS-7 — Remote access authorization and monitoring": "PS-7 — Autorisation et surveillance de l'accès à distance",
 "Remote-access policy acknowledgments; access logs showing session monitoring": "Accusés de réception de la politique d'accès à distance; journaux d'accès montrant la surveillance des sessions",
 "PS-9 — Session lock after inactivity": "PS-9 — Verrouillage de session après inactivité",
 "GPO / MDM config screenshots; inactivity-lock test results": "Captures d'écran de configuration GPO/GAM; résultats des tests de verrouillage par inactivité",
 "PS-10 — Session termination after defined period": "PS-10 — Fin de session après une période définie",
 "System configuration records for session termination settings": "Registres de configuration système pour les paramètres de fin de session",
 "PS-11 — Wireless access controls and authorization": "PS-11 — Contrôles et autorisation de l'accès sans fil",
 "Wireless network config; WPA3/enterprise auth records; wireless access logs": "Configuration du réseau sans fil; registres d'authentification WPA3/entreprise; journaux d'accès sans fil",
 "PS-12, PS-13 — Mobile device policy; MDM enrollment; CUI controls on mobile": "PS-12, PS-13 — Politique des appareils mobiles; inscription à la GAM; contrôles des CUI sur mobile",
 "MDM enrollment reports; mobile device policy acknowledgments": "Rapports d'inscription à la GAM; accusés de réception de la politique des appareils mobiles",
 "PS-14 — External system connection controls; use agreements": "PS-14 — Contrôles de connexion des systèmes externes; ententes d'utilisation",
 "External-connection approval records; signed use agreements; CUI-on-external-system inventory": "Registres d'approbation de connexion externe; ententes d'utilisation signées; inventaire des CUI sur systèmes externes",
 "Date": "Date", "Author / Role": "Auteur / Rôle", "Summary of Changes": "Résumé des modifications",
 "Initial policy release.": "Publication initiale de la politique.",
 "[INSERT ORGANIZATION LOGO]": "[INSÉRER LE LOGO DE L'ORGANISATION]",
}

def set_text(p, txt):
    if not p.runs:
        p.add_run(txt); return
    p.runs[0].text = txt
    for r in p.runs[1:]:
        r.text = ""

def translate_paras(paras, stats):
    for p in paras:
        t = p.text.strip()
        if not t: continue
        if t in TRANS and TRANS[t]:
            set_text(p, TRANS[t]); stats[0]+=1
        else:
            stats[1]+=1
            if t not in stats[2]: stats[2].append(t)

d = Document("policy_library/Access_Control_Policy.docx")
stats=[0,0,[]]
translate_paras(d.paragraphs, stats)
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            translate_paras(cell.paragraphs, stats)
d.save("policy_library/Access_Control_Policy.fr.docx")
print(f"translated={stats[0]} untranslated={stats[1]}")
for u in stats[2][:40]:
    print("  UNTRANSLATED:", repr(u[:90]))
