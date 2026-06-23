import sys; sys.path.insert(0,'tools')
import policy_builder as B

d = B.new_doc()
ORG = "[Organization Name]"
B.title_block(d, ORG, "Audit & Accountability Policy")

B.table(d, [
    ["Field","Value","Field","Value"],
    ["Policy Name","Audit & Accountability Policy","Version","[1.0]"],
    ["Policy Owner","[IT Manager / ISSO]","Approved By","[CEO / CISO]"],
    ["Effective Date","[YYYY-MM-DD]","Last Reviewed","[YYYY-MM-DD]"],
    ["Review Cadence","Annual (or upon major change)","Classification","[CUI / Internal]"],
])

B.heading(d, "1.  Purpose")
B.body(d, "This policy establishes how [Organization Name] generates, protects, reviews, and retains audit records (logs) for systems that process, store, or transmit Controlled Unclassified Information (CUI). Reliable, tamper-resistant logs are essential to detect security incidents, hold individuals accountable for their actions, support investigations, and demonstrate compliance. By defining what is logged, how logs are protected, and how they are reviewed, this policy ensures that significant events can be reconstructed and that malicious or unauthorized activity does not go unnoticed.")

B.heading(d, "2.  Scope")
B.body(d, "This policy applies to:")
B.bullet(d, "All information systems owned, operated, or managed by [Organization Name] that process, store, or transmit CUI, including servers, endpoints, network devices, cloud services, and applications.")
B.bullet(d, "All system and security components capable of generating audit records, including identity providers, firewalls, EDR, and SaaS platforms.")
B.bullet(d, "All personnel responsible for configuring, operating, reviewing, or protecting audit logs.")
B.body(d, "Where CUI is involved, the logging requirements in this policy are mandatory and supersede convenience or storage-cost preferences.")

B.heading(d, "3.  Definitions")
B.body(d, "Audit Record (Log): A time-stamped record of an event on a system (e.g., a logon, a configuration change, access to a file).")
B.body(d, "Auditable Event: An event the organization has decided is security-relevant and must be logged.")
B.body(d, "Log Source: Any system or component that produces audit records.")
B.body(d, "SIEM (Security Information and Event Management): A central platform that aggregates, correlates, and alerts on logs from many sources.")
B.body(d, "Time Synchronization: Keeping system clocks aligned to an authoritative time source so events across systems can be correlated.")
B.body(d, "Log Integrity: Assurance that audit records have not been altered or deleted since they were created.")

B.heading(d, "4.  Roles & Responsibilities")
B.table(d, [
    ["Role","Key Responsibilities"],
    ["[IT Manager / ISSO]","Own this policy; define auditable events; ensure logs are reviewed and retained; report findings to leadership."],
    ["[System Administrator]","Enable and configure logging on systems; forward logs to the central store; maintain time synchronization; respond to logging failures."],
    ["[Security Analyst / Reviewer]","Review and analyze logs and alerts; investigate anomalies; document findings and escalate incidents."],
    ["All Users","Do not disable logging or tamper with logs; report suspected log tampering or security events."],
    ["[Senior Leadership / CISO]","Approve this policy; ensure adequate logging tools and storage; accept residual risk for documented exceptions."],
])

B.heading(d, "5.  Policy Statements")
B.ps(d,"PS-1  Define Auditable Events","[Organization Name] must define, document, and maintain the list of auditable events for each CUI system. At a minimum this includes successful and failed logons, account and privilege changes, access to CUI, security configuration changes, use of privileged functions, and security tool alerts. The list must be reviewed [organization-defined: at least annually; recommended: annually].")
B.ps(d,"PS-2  Audit Record Content","Each audit record must capture sufficient detail to reconstruct the event: what type of event occurred, when it occurred (date and time), where it occurred (system/source), the source of the event (user or process), and the outcome (success/failure). Additional context (e.g., affected resource) must be captured where available.")
B.ps(d,"PS-3  Audit Record Generation","Logging must be enabled on all in-scope systems and components for the defined auditable events. Logging must be turned on by default in system baselines and must not be disabled without documented approval from the policy owner.")
B.ps(d,"PS-4  Time Synchronization & Timestamps","All in-scope systems must synchronize their clocks to an authoritative internal time source that is itself synchronized to a trusted external source (e.g., an NTP hierarchy). Audit records must use accurate timestamps recorded to at least the second, in a consistent time zone (UTC recommended), so events can be correlated across systems.")
B.ps(d,"PS-5  Centralized Collection","Security-relevant logs must be forwarded from their sources to a central, time-synchronized log store or SIEM. Centralization protects logs from local tampering, enables correlation across systems, and supports retention.")
B.ps(d,"PS-6  Protection of Audit Information","Audit records and audit tools must be protected from unauthorized access, modification, and deletion. Access to logs must be restricted to authorized roles on a need-to-know basis, write/delete access must be tightly limited, and integrity controls (e.g., write-once storage, hashing, or access logging on the log store) must be applied.")
B.ps(d,"PS-7  Logging Failure Response","Systems must alert designated personnel when a logging process fails (e.g., a log source stops reporting or storage is exhausted). The organization must define the response — at a minimum, alert and investigate — and must take action to restore logging promptly.")
B.ps(d,"PS-8  Audit Storage Capacity","Sufficient storage must be allocated for audit records to meet the retention requirement, and storage utilization must be monitored so that logging is not interrupted by a full log store.")
B.ps(d,"PS-9  Review, Analysis & Reporting","Logs and alerts must be reviewed and analyzed [organization-defined: at least weekly, and on alert; recommended: continuous alerting with weekly review] for indications of inappropriate or unusual activity. Findings must be documented, and confirmed security events must be handled under the Incident Response Policy.")
B.ps(d,"PS-10  Alerting & Correlation","Where a SIEM or equivalent is available, correlation rules and alerts must be configured for high-risk patterns (e.g., repeated authentication failures, privilege escalation, mass file access or deletion, disabled security tools). Alerts must route to a monitored queue.")
B.ps(d,"PS-11  Log Retention","Audit records must be retained for [organization-defined: at least 12 months online and 90 days immediately available; recommended: 12 months, aligned to contractual/CMMC requirements] to support after-the-fact investigations, then disposed of securely.")
B.ps(d,"PS-12  Privileged Activity Logging","All actions taken by privileged/administrator accounts must be logged, including the use of privileged functions, and reviewed with particular attention given their elevated risk.")

B.heading(d, "6.  Procedures & Audit Evidence")
B.body(d, "6.1  Logging Configuration")
B.bullet(d, "Define auditable events and document them in the logging standard; encode them into system baselines (GPO/MDM/config).")
B.bullet(d, "Enable logging on each in-scope system and forward logs to the central store/SIEM.")
B.bullet(d, "Configure time synchronization to the approved time source on every system.")
B.body(d, "Audit evidence: Logging standard / auditable-event list (dated); baseline configurations showing logging enabled; SIEM source inventory; NTP configuration.")
B.body(d, "6.2  Review & Response")
B.bullet(d, "Reviewers analyze alerts continuously and perform a documented log review [weekly].")
B.bullet(d, "Logging failures generate alerts; administrators restore logging and record the cause and resolution.")
B.bullet(d, "Confirmed security events are escalated to incident response.")
B.body(d, "Audit evidence: Dated log-review records; alert/triage tickets; logging-failure incident records; correlation rule list.")
B.body(d, "6.3  Protection & Retention")
B.bullet(d, "Restrict and log access to the log store; apply integrity protection (write-once / hashing).")
B.bullet(d, "Monitor log-storage capacity; retain records for the defined period; dispose securely at end of life.")
B.body(d, "Audit evidence: Log-store access-control configuration and access logs; retention configuration; storage-capacity monitoring; disposal records.")

B.heading(d, "7.  Compliance, Monitoring & Enforcement")
B.body(d, "The [IT Manager / ISSO] shall review compliance with this policy at least [quarterly] by verifying that in-scope systems are logging the defined events, that reviews are occurring, and that retention and protection controls are operating. Results shall be reported to [Senior Leadership / CISO].")
B.body(d, "Violations may result in:")
B.bullet(d, "Remediation of misconfigured or disabled logging within a defined timeframe.")
B.bullet(d, "Disciplinary action up to and including termination for deliberately disabling or tampering with logs, consistent with [Organization Name]'s disciplinary procedures.")
B.bullet(d, "Referral to law enforcement and notification to government contracting officers where CUI or criminal activity is involved, as required.")

B.heading(d, "8.  Exceptions")
B.body(d, "Exceptions to this policy must be rare. To request one:")
B.bullet(d, "The requester submits a written exception to the [IT Manager / ISSO] describing the requirement that cannot be met, the business justification, the risk, and compensating controls.")
B.bullet(d, "The [IT Manager / ISSO] assesses the risk and recommends; [Senior Leadership / CISO] approves or denies in writing.")
B.bullet(d, "Approved exceptions are logged in the Exception Register with an expiry of no more than [12 months] and reflected as a POA&M item in the SSP where CUI systems are affected.")

B.heading(d, "9.  Related Documents")
B.bullet(d, "System Security Plan (SSP) — system inventory, CUI categories, and control implementation details.")
B.bullet(d, "Incident Response Policy — handling of confirmed security events detected through log review.")
B.bullet(d, "Access Control Policy — account and privileged-access controls that audit logging records.")
B.bullet(d, "Configuration Management Policy — baselines that enable logging and time synchronization.")
B.bullet(d, "System & Information Integrity Policy — monitoring and alerting that consume audit data.")

B.heading(d, "10.  Control Mapping")
B.body(d, "The table below maps each required control to the policy statement(s) that satisfy it and the evidence an auditor will seek. Control IDs are from NIST SP 800-171 Rev 3 (family 03.03); they are equally applicable to CMMC 2.0 Level 2 and ITSP.10.171 (CPCSC).")
B.table(d, [
    ["Control ID","Policy Statement(s)","Expected Audit Evidence"],
    ["03.03.01  Event Logging","PS-1, PS-3 — defined auditable events; logging enabled","Auditable-event list; baseline configs; SIEM source inventory"],
    ["03.03.02  Audit Record Content","PS-2 — record content (who/what/when/where/outcome)","Sample log records showing required fields"],
    ["03.03.03  Audit Record Generation","PS-3, PS-12 — generation on all sources; privileged activity","Logging configuration; privileged-action logs"],
    ["03.03.04  Response to Audit Logging Process Failures","PS-7, PS-8 — failure alerting and response; capacity","Logging-failure alerts and tickets; storage-capacity monitoring"],
    ["03.03.05  Audit Record Review, Analysis & Reporting","PS-9, PS-10 — review, analysis, alerting, correlation","Dated log-review records; alert/triage tickets; correlation rules"],
    ["03.03.06  Audit Record Reduction & Report Generation","PS-5, PS-10 — centralized SIEM; reporting","SIEM reports; search/report capability evidence"],
    ["03.03.07  Time Stamps","PS-4 — synchronized clocks; accurate timestamps","NTP configuration; sample timestamped records"],
    ["03.03.08  Protection of Audit Information","PS-6, PS-11 — protection, access control, retention","Log-store access controls and access logs; integrity/retention config"],
])

B.heading(d, "11.  Revision History")
B.table(d, [
    ["Version","Date","Author / Role","Summary of Changes"],
    ["1.0","[YYYY-MM-DD]","[IT Manager / ISSO]","Initial policy release."],
])

d.save("policy_library/Audit_Accountability_Policy.docx")
# refresh FR placeholder removed (will re-translate after EN approved)
print("Audit_Accountability_Policy.docx rebuilt")
from docx import Document
dd=Document("policy_library/Audit_Accountability_Policy.docx")
print("paragraphs:", len([p for p in dd.paragraphs if p.text.strip()]), "tables:", len(dd.tables))
